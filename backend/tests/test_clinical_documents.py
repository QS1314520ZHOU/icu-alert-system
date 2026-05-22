from __future__ import annotations

import io
import sys
from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from app.clinical_documents.document_generator import ProgressNoteGenerator, render_note_preview
from app.clinical_documents.exporter import export_progress_note_docx
from app.clinical_documents.schemas import (
    AlertItem,
    Basics,
    DrugEvent,
    LabDelta,
    ProgressNoteContext,
    RoundingWorkbenchDraft,
    Scores,
    Ventilator,
    VitalStat,
    Vitals,
)


def make_context(with_vent: bool = True) -> ProgressNoteContext:
    return ProgressNoteContext(
        patient_id="P001",
        window_start="2026-05-21 13:45",
        window_end="2026-05-22 13:45",
        basics=Basics(name="王小明", bed="03", age=39, sex="男", day=31, diagnosis="重症肺炎/ARDS风险"),
        v=Vitals(
            hr=VitalStat(min=75, max=99, trend="平稳"),
            map=VitalStat(min=67, max=82, trend="波动"),
            spo2=VitalStat(min=89, max=99, trend="下降"),
            temp=VitalStat(min=36.8, max=38.2, trend="上升"),
            rr=VitalStat(min=18, max=28, trend="上升"),
            events=[],
        ),
        labs=[LabDelta(id=1, name="乳酸", prev=2.1, curr=3.5, unit="mmol/L", flag="high")],
        drugs=[DrugEvent(id=1, time_hm="10:30", action="新增", name="哌拉西林他唑巴坦", dose_after="4.5g")],
        vent=Ventilator(mode="PRVC", fio2=0.6, peep=10, vt=420, pplat=24, pf_ratio=150) if with_vent else None,
        alerts=[AlertItem(id=1, type="ards", severity="high", count=2, active=True)],
        scores=Scores(gcs=12, sofa=6, apache=18),
    )


def test_workbench_schema_and_fallback_draft() -> None:
    ctx = make_context()
    gen = ProgressNoteGenerator(None)
    citations = gen._build_citations(ctx)
    draft = gen._build_skeleton(ctx, citations)
    draft = gen._validate_and_repair_workbench(draft, citations)
    draft["quality_checks"] = gen._run_quality_checks(draft, ctx)
    draft["note_preview"] = render_note_preview(draft)

    parsed = RoundingWorkbenchDraft(**draft)

    assert parsed.schema_version == "icu_rounding_workbench.v1"
    assert parsed.content_type == "rounding_workbench"
    assert parsed.patient_banner.bed_no == "03"
    assert len(parsed.system_ap) >= 10
    assert any(card.system == "resp" for card in parsed.system_ap)
    assert parsed.note_preview.generated_text.startswith("A/P")


def test_citation_integrity_for_workbench() -> None:
    ctx = make_context()
    gen = ProgressNoteGenerator(None)
    citations = gen._build_citations(ctx)
    draft = gen._build_skeleton(ctx, citations)
    valid_refs = {item["id"] for item in citations}

    for card in draft["system_ap"]:
        for group in ("status", "trend", "assessment", "plan_items"):
            for statement in card[group]:
                for ref in statement["evidence_refs"]:
                    assert ref in valid_refs
                if statement["kind"] == "fact":
                    assert statement["evidence_refs"]
                if statement["kind"] == "recommendation" and not statement["evidence_refs"]:
                    assert statement["missing_data"] or statement["review_required"]


def test_note_preview_is_derived_from_structured_plan() -> None:
    ctx = make_context()
    gen = ProgressNoteGenerator(None)
    draft = gen._build_skeleton(ctx, gen._build_citations(ctx))
    original = render_note_preview(draft)["generated_text"]

    resp = next(card for card in draft["system_ap"] if card["system"] == "resp")
    resp["plan_items"].append({
        "id": "resp_plan_abg",
        "kind": "recommendation",
        "text": "16:00复查血气并记录P/F ratio。",
        "evidence_refs": ["VT0"],
        "missing_data": [],
        "review_required": True,
    })
    updated = render_note_preview(draft)["generated_text"]

    assert original != updated
    assert "16:00复查血气" in updated


def test_missing_data_quality_checks_without_ventilator() -> None:
    ctx = make_context(with_vent=False)
    gen = ProgressNoteGenerator(None)
    draft = gen._build_skeleton(ctx, gen._build_citations(ctx))
    checks = gen._run_quality_checks(draft, ctx)

    assert "FiO2" in checks["critical_missing_data"]
    assert "PEEP" in checks["critical_missing_data"]
    assert any("呼吸机参数" in warning for warning in checks["warnings"])


def test_export_workbench_docx_contains_chinese_sections() -> None:
    ctx = make_context()
    gen = ProgressNoteGenerator(None)
    citations = gen._build_citations(ctx)
    content = gen._build_skeleton(ctx, citations)
    content["quality_checks"] = gen._run_quality_checks(content, ctx)
    content["note_preview"] = render_note_preview(content)
    doc_bytes = export_progress_note_docx({
        "patient_id": "P001",
        "status": "finalized",
        "finalized_by": "张医生",
        "finalized_at": "2026-05-22T13:50:00",
        "current_content": content,
    })

    parsed = Document(io.BytesIO(doc_bytes))
    text = "\n".join(p.text for p in parsed.paragraphs)

    assert "ICU 24小时病程记录" in text
    assert "Assessment & Plan" in text
    assert "今日目标" in text
    assert "医师签名" in text


def test_export_legacy_soap_docx_still_works() -> None:
    doc_bytes = export_progress_note_docx({
        "patient_id": "P002",
        "status": "draft",
        "current_content": {
            "subjective": "患者主诉需床旁补充。",
            "objective": {"vitals": "HR 80, MAP 75"},
            "assessment": {"呼吸": "氧合需复核"},
            "plan": ["复查血气"],
            "overall_trend": "平稳",
            "key_concerns": ["低氧风险"],
        },
        "context_snapshot": {"basics": {"bed": "05", "age": 55, "sex": "男", "day": 5, "diagnosis": "重症肺炎"}},
    })

    parsed = Document(io.BytesIO(doc_bytes))
    text = "\n".join(p.text for p in parsed.paragraphs)
    assert "主观症状" in text
    assert "复查血气" in text


@pytest.mark.asyncio
async def test_update_draft_rerenders_note_preview() -> None:
    from app.routers.clinical_documents import UpdateDraftRequest, update_draft

    ctx = make_context()
    gen = ProgressNoteGenerator(None)
    content = gen._build_skeleton(ctx, gen._build_citations(ctx))
    content["system_ap"][0]["plan_items"].append({
        "id": "neuro_plan_test",
        "kind": "recommendation",
        "text": "补充神经查体。",
        "evidence_refs": [],
        "missing_data": ["床旁查体"],
        "review_required": True,
    })

    class FakeCollection:
        def __init__(self):
            self.updated = None

        async def find_one(self, query):
            return {"_id": "draft_123", "status": "draft", "current_content": content}

        async def count_documents(self, query):
            return 0

        async def insert_one(self, doc):
            return None

        async def update_one(self, query, payload):
            self.updated = payload["$set"]["current_content"]

    class FakeDb:
        def __init__(self):
            self.main = FakeCollection()
            self.ver = FakeCollection()

        def col(self, name):
            return self.ver if name == "clinical_document_versions" else self.main

    db = FakeDb()
    result = await update_draft("draft_123", UpdateDraftRequest(content=content), db)

    assert result["ok"] is True
    assert db.main.updated["note_preview"]["generated_from_hash"]


@pytest.mark.asyncio
async def test_context_builder_queries() -> None:
    from app.clinical_documents.context_builder import ProgressNoteContextBuilder

    class FakeCursor:
        def __init__(self, items):
            self.items = items
            self.index = 0

        def sort(self, *args, **kwargs):
            return self

        async def to_list(self, length):
            return self.items

        def __aiter__(self):
            return self

        async def __anext__(self):
            if self.index < len(self.items):
                item = self.items[self.index]
                self.index += 1
                return item
            raise StopAsyncIteration

    class FakeCollection:
        def __init__(self, data=None):
            self.data = data or []

        async def find_one(self, query, sort=None):
            return self.data[0] if self.data else None

        def find(self, query):
            return FakeCursor(self.data)

        async def aggregate(self, pipeline):
            return FakeCursor([])

    class FakeDb:
        def __init__(self):
            now = datetime.now()
            self.collections = {
                "patient": FakeCollection([{"_id": "P001", "hisPid": "HIS999", "bedNo": "05", "age": 55, "sex": "男", "diagnosis": "重症肺炎", "admissionTime": now}]),
                "bGATemp": FakeCollection([{"mrn": "HIS999", "inputTime": now, "ventMode": "SIMV", "param_HR": 85, "param_nibp_m": 90, "param_spo2": 98, "param_T": 37.2, "param_resp": 18, "bedsides": [{"code": "param_FiO2", "fVal": 40}, {"code": "param_vent_measure_peep", "fVal": 8}, {"code": "param_vent_vt", "fVal": 420}, {"code": "param_vent_plat_pressure", "fVal": 21}, {"code": "param_bg_P/Fratio", "fVal": 243}]}]),
                "drugExe": FakeCollection([{"pid": "P001", "drugName": "去甲肾上腺素", "executeTime": now, "dose": 0.1, "doseUnit": "ug/kg/min"}]),
                "score": FakeCollection([{"patient_id": "P001", "score_type": "GCS", "score": 15, "calc_time": now}]),
                "alert_records": FakeCollection([]),
                "VI_ICU_EXAM_ITEM": FakeCollection([{"hisPid": "HIS999", "itemName": "乳酸", "result": "2.5", "refHigh": 2.0, "refLow": 0.5, "authTime": now}]),
                "VI_ICU_ZYYZ": FakeCollection([]),
            }

        def col(self, name):
            return self.collections.get(name, FakeCollection())

        def dc_col(self, name):
            return self.collections.get(name, FakeCollection())

    ctx = await ProgressNoteContextBuilder(FakeDb()).build("P001", hours=24)

    assert ctx.basics.bed == "05"
    assert ctx.basics.age == 55
    assert ctx.vent is not None
    assert ctx.vent.pf_ratio == 243
