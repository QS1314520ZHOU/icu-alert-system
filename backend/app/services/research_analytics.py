from __future__ import annotations

import asyncio
import io
import logging
import math
import re
import uuid
import zipfile
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any

try:
    import pandas as pd
except Exception:  # pragma: no cover
    pd = None  # type: ignore
from bson import ObjectId
try:
    from scipy import stats
except Exception:  # pragma: no cover
    stats = None  # type: ignore
try:
    from sklearn import metrics
except Exception:  # pragma: no cover
    metrics = None  # type: ignore
try:
    from sklearn.linear_model import LinearRegression, LogisticRegression
except Exception:  # pragma: no cover
    LinearRegression = None  # type: ignore
    LogisticRegression = None  # type: ignore

try:
    import numpy as np
except Exception:  # pragma: no cover
    np = None  # type: ignore

try:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
except Exception:  # pragma: no cover
    matplotlib = None  # type: ignore
    plt = None  # type: ignore

try:
    import seaborn as sns
except Exception:  # pragma: no cover
    sns = None  # type: ignore

try:
    import statsmodels.api as sm
    import statsmodels.formula.api as smf
    from statsmodels.stats.anova import anova_lm
except Exception:  # pragma: no cover
    sm = None  # type: ignore
    smf = None  # type: ignore
    anova_lm = None  # type: ignore

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None  # type: ignore

try:
    from lifelines import CoxPHFitter, KaplanMeierFitter
    from lifelines.statistics import logrank_test, multivariate_logrank_test
except Exception:  # pragma: no cover
    CoxPHFitter = None  # type: ignore
    KaplanMeierFitter = None  # type: ignore
    logrank_test = None  # type: ignore
    multivariate_logrank_test = None  # type: ignore

from app.utils.serialization import safe_oid, serialize_doc
from app.utils.patient_helpers import research_patient_scope_query

logger = logging.getLogger("icu-alert")

RESEARCH_EXPORT_DIR = Path("backend/exports/research_analytics")
RESEARCH_EXPORT_DIR.mkdir(parents=True, exist_ok=True)

FIGURE_WIDTH_CM_SINGLE = 8.5
FIGURE_WIDTH_CM_DOUBLE = 17.5


def _require_research_analytics_deps() -> None:
    missing: list[str] = []
    if pd is None:
        missing.append("pandas")
    if missing:
        raise RuntimeError(f"科研分析依赖缺失: {', '.join(missing)}")

VITAL_CODE_MAP: dict[str, list[str]] = {
    "hr": ["param_HR", "param_PR"],
    "map": ["param_ibp_m", "param_nibp_m"],
    "sbp": ["param_ibp_s", "param_nibp_s", "param_abp_s"],
    "dbp": ["param_ibp_d", "param_nibp_d", "param_abp_d"],
    "spo2": ["param_spo2"],
    "rr": ["param_resp"],
    "temp": ["param_T", "param_temp"],
    "temperature": ["param_T", "param_temp"],
}

LAB_NAME_MAP: dict[str, list[str]] = {
    "lactate": ["lactate", "lac", "乳酸", "blood lactate", "动脉血乳酸"],
    "creatinine": ["creatinine", "cr", "肌酐", "scr"],
    "wbc": ["wbc", "白细胞", "white blood cell", "leukocyte", "白细胞计数"],
    "pct": ["pct", "降钙素原", "procalcitonin"],
    "albumin": ["albumin", "alb", "白蛋白"],
    "hemoglobin": ["hemoglobin", "hb", "hgb", "血红蛋白"],
    "platelet": ["platelet", "plt", "血小板", "血小板计数", "血小板总数"],
    "pf_ratio": ["pf ratio", "p/f", "氧合指数", "pao2/fio2", "P/F Ratio"],
    "bnp": ["bnp", "brain natriuretic peptide", "脑钠肽", "nt-probnp", "proBNP", "B型钠尿肽前体"],
}

TREND_LAB_FIELD_MAP: dict[str, str] = {
    "lactate": "lactate",
    "lactate_admission": "lactate",
    "creatinine": "creatinine",
    "creatinine_admission": "creatinine",
    "albumin": "albumin",
    "albumin_admission": "albumin",
    "pct": "pct",
    "pct_admission": "pct",
    "wbc": "wbc",
    "wbc_admission": "wbc",
    "hemoglobin": "hemoglobin",
    "hemoglobin_admission": "hemoglobin",
    "platelet": "platelet",
    "platelet_admission": "platelet",
    "pf_ratio": "pf_ratio",
    "pf_ratio_admission": "pf_ratio",
    "bnp": "bnp",
    "bnp_admission": "bnp",
}


def _get_research_cfg(config) -> dict[str, Any]:
    cfg = (getattr(config, "yaml_cfg", None) or {}) if config is not None else {}
    value = cfg.get("research", {}) if isinstance(cfg, dict) else {}
    return value if isinstance(value, dict) else {}


def _p_display(p_value: float | None) -> str:
    if p_value is None or not isinstance(p_value, (int, float)) or math.isnan(float(p_value)):
        return "—"
    if p_value < 0.001:
        return "<0.001"
    return f"{p_value:.3f}"


def _as_number(value: Any) -> float | None:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return 1.0 if value else 0.0
    if isinstance(value, (int, float)):
        num = float(value)
        return num if math.isfinite(num) else None
    text = str(value).strip()
    if not text or text.lower() in {"nan", "none", "null"}:
        return None
    try:
        num = float(text)
        return num if math.isfinite(num) else None
    except Exception:
        return None


def _coerce_binary(value: Any) -> int | None:
    if value is None or value == "":
        return None
    if isinstance(value, bool):
        return 1 if value else 0
    if isinstance(value, (int, float)):
        return 1 if float(value) > 0 else 0
    text = str(value).strip().lower()
    if text in {"1", "yes", "y", "true", "male", "m", "dead", "death", "死亡", "有"}:
        return 1
    if text in {"0", "no", "n", "false", "female", "f", "alive", "存活", "无"}:
        return 0
    return None


def _normalize_dt(value: datetime) -> datetime:
    if value.tzinfo is None:
        return value.replace(tzinfo=timezone.utc)
    return value.astimezone(timezone.utc)


def _parse_dt(value: Any) -> datetime | None:
    if isinstance(value, datetime):
        return _normalize_dt(value)
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    text = text.replace("Z", "+00:00")
    try:
        return _normalize_dt(datetime.fromisoformat(text))
    except Exception:
        return None


def _to_series_numeric(series: pd.Series) -> Any:
    return pd.to_numeric(series, errors="coerce").to_numpy(dtype=float)


def _ensure_patient_id_set(patient_ids: list[str] | None) -> list[str]:
    values: list[str] = []
    for raw in patient_ids or []:
        text = str(raw or "").strip()
        if text:
            values.append(text)
    dedup: list[str] = []
    seen: set[str] = set()
    for item in values:
        if item in seen:
            continue
        seen.add(item)
        dedup.append(item)
    return dedup


def _infer_outcome(doc: dict[str, Any]) -> str:
    if any(doc.get(key) for key in ("deathTime", "deceasedAt", "死亡时间")):
        return "dead"
    for key in ("icu_mortality", "hospital_mortality", "mortality", "isDeath", "death"):
        result = _coerce_binary(doc.get(key))
        if result is not None:
            return "dead" if result else "alive"
    for raw in (
        doc.get("outcome"),
        doc.get("status"),
        doc.get("dischargeOutcome"),
        doc.get("leaveType"),
        doc.get("dischargeType"),
        doc.get("dischargeDisposition"),
        doc.get("remark"),
        doc.get("deathReason"),
    ):
        text = str(raw or "").strip().lower()
        if any(token in text for token in ["dead", "death", "deceased", "死亡", "抢救无效", "expired"]):
            return "dead"
        if any(token in text for token in ["alive", "survive", "存活", "转出", "discharged", "出院", "好转"]):
            return "alive"
    return "alive"


def _infer_los_days(doc: dict[str, Any]) -> float | None:
    for key in ("los_icu_days", "icu_los_days", "losDays", "icuStayDays"):
        value = _as_number(doc.get(key))
        if value is not None:
            return value
    start = _parse_dt(doc.get("icuAdmissionTime") or doc.get("admissionTime") or doc.get("inTime"))
    end = _parse_dt(doc.get("icuDischargeTime") or doc.get("dischargeTime") or doc.get("outTime"))
    if start and end and end >= start:
        return (end - start).total_seconds() / 86400.0
    return None


def _infer_discharge_dest(doc: dict[str, Any]) -> str:
    text = str(
        doc.get("discharge_dest")
        or doc.get("dischargeDestination")
        or doc.get("transfer_dest")
        or doc.get("outcome")
        or doc.get("status")
        or ""
    ).strip().lower()
    if not text:
        return "未知"
    if any(token in text for token in ["dead", "death", "死亡"]):
        return "ICU内死亡"
    if any(token in text for token in ["auto", "自动出院"]):
        return "自动出院"
    if any(token in text for token in ["转科", "transfer", "ward"]):
        return "转科"
    if any(token in text for token in ["alive", "survive", "存活", "出科", "discharge"]):
        return "存活出科"
    return str(doc.get("discharge_dest") or doc.get("outcome") or "未知")


def _infer_los_group(los_days: float | None) -> str:
    if los_days is None:
        return "未知"
    if los_days < 3:
        return "<3天"
    if los_days < 7:
        return "3-7天"
    return ">=7天"


def _is_normal_distribution(values: Any) -> bool:
    if np is None:
        return False
    arr = np.asarray(values, dtype=float)
    arr = arr[np.isfinite(arr)]
    if len(arr) < 3:
        return False
    if len(arr) > 5000:
        arr = np.random.choice(arr, 5000, replace=False)
    try:
        _, p_value = stats.shapiro(arr)
        return bool(p_value > 0.05)
    except Exception:
        return False


def _format_continuous(values: Any, normal: bool) -> str:
    if np is None:
        return "—"
    arr = np.asarray(values, dtype=float)
    arr = arr[np.isfinite(arr)]
    if len(arr) == 0:
        return "—"
    if normal:
        mean = float(np.mean(arr))
        std = float(np.std(arr, ddof=1)) if len(arr) > 1 else 0.0
        return f"{mean:.2f} ± {std:.2f}"
    q1, q2, q3 = np.percentile(arr, [25, 50, 75])
    return f"{q2:.2f} ({q1:.2f}-{q3:.2f})"


def _format_count_ratio(count: int, total: int) -> str:
    if total <= 0:
        return "0 (0.0%)"
    return f"{count} ({count * 100.0 / total:.1f}%)"


def _chi_square_or_fisher(table: Any) -> tuple[str, float | None]:
    if np is None:
        return "χ²=—", None
    arr = np.asarray(table, dtype=float)
    if arr.size == 0:
        return "χ²=—", None
    try:
        if arr.shape == (2, 2):
            chi2, p_value, _, expected = stats.chi2_contingency(arr)
            if np.any(expected < 5):
                _, p_fisher = stats.fisher_exact(arr)
                return "Fisher", float(p_fisher)
            return f"χ²={chi2:.2f}", float(p_value)
        chi2, p_value, _, _ = stats.chi2_contingency(arr)
        return f"χ²={chi2:.2f}", float(p_value)
    except Exception:
        return "χ²=—", None


def _apply_filter(df: pd.DataFrame, rule: dict[str, Any]) -> pd.Series:
    if df.empty:
        return pd.Series([], dtype=bool)
    if not rule:
        return pd.Series([True] * len(df), index=df.index)
    mask = pd.Series([True] * len(df), index=df.index)
    for key, cond in rule.items():
        if key in {"$and", "$or"} and isinstance(cond, list):
            sub_masks = [_apply_filter(df, item if isinstance(item, dict) else {}) for item in cond]
            if not sub_masks:
                continue
            combined = sub_masks[0]
            for sub in sub_masks[1:]:
                combined = (combined & sub) if key == "$and" else (combined | sub)
            mask = mask & combined
            continue
        if key == "diagnosis_contains":
            series = df.get("primary_diagnosis", pd.Series([""] * len(df), index=df.index)).astype(str)
            mask = mask & series.str.contains(str(cond or ""), case=False, na=False)
            continue
        if key == "diagnosis_not_contains":
            series = df.get("primary_diagnosis", pd.Series([""] * len(df), index=df.index)).astype(str)
            mask = mask & (~series.str.contains(str(cond or ""), case=False, na=False))
            continue
        if key not in df.columns:
            mask = mask & False
            continue
        col = df[key]
        if isinstance(cond, dict):
            for op, value in cond.items():
                if op == "$lt":
                    mask = mask & (col < value)
                elif op == "$lte":
                    mask = mask & (col <= value)
                elif op == "$gt":
                    mask = mask & (col > value)
                elif op == "$gte":
                    mask = mask & (col >= value)
                elif op == "$in":
                    mask = mask & col.isin(value if isinstance(value, list) else [value])
                elif op == "$nin":
                    mask = mask & (~col.isin(value if isinstance(value, list) else [value]))
                elif op == "$regex":
                    mask = mask & col.astype(str).str.contains(str(value), case=False, na=False)
        else:
            mask = mask & (col == cond)
    return mask


async def _load_patient_dataframe(
    patient_ids: list[str] | None,
    db,
    *,
    max_patients: int = 10000,
) -> pd.DataFrame:
    pid_values = _ensure_patient_id_set(patient_ids)
    oid_values: list[ObjectId] = []
    string_id_values: list[str] = []
    his_pid_values: list[str] = []
    for token in pid_values:
        oid = safe_oid(token)
        if oid is not None:
            oid_values.append(oid)
        else:
            string_id_values.append(token)
            his_pid_values.append(token)

    query: dict[str, Any]
    if oid_values or string_id_values or his_pid_values:
        or_terms: list[dict[str, Any]] = []
        if oid_values:
            or_terms.append({"_id": {"$in": oid_values}})
        if string_id_values:
            # 某些环境 patient._id 为字符串（非 ObjectId）；同时尝试按字符串 _id 命中。
            or_terms.append({"_id": {"$in": string_id_values}})
        if his_pid_values:
            or_terms.append({"$or": [{"hisPid": {"$in": his_pid_values}}, {"hisPID": {"$in": his_pid_values}}]})
        query = {"$or": or_terms}
    else:
        query = {}

    projection = {
        "_id": 1,
        "hisPid": 1,
        "hisPID": 1,
        "hisPatientId": 1,
        "patientId": 1,
        "patientID": 1,
        "pid": 1,
        "mrn": 1,
        "hisMrn": 1,
        "name": 1,
        "sex": 1,
        "gender": 1,
        "age": 1,
        "birthday": 1,
        "dept": 1,
        "hisDept": 1,
        "deptCode": 1,
        "status": 1,
        "outcome": 1,
        "deathTime": 1,
        "deceasedAt": 1,
        "死亡时间": 1,
        "admissionTime": 1,
        "icuAdmissionTime": 1,
        "dischargeTime": 1,
        "icuDischargeTime": 1,
        "diagnosis": 1,
        "clinicalDiagnosis": 1,
        "admissionDiagnosis": 1,
        "mechanical_ventilation": 1,
        "crrt": 1,
        "vasopressor": 1,
        "hospital_mortality": 1,
        "mortality_28d": 1,
        "mortality28d": 1,
        "discharge_dest": 1,
        "dischargeDestination": 1,
        "transfer_dest": 1,
        "sofa_admission": 1,
        "apache2": 1,
        "los_icu_days": 1,
    }
    cursor = db.col("patient").find(query, projection).limit(max_patients)
    docs = [serialize_doc(doc) async for doc in cursor]
    if not docs:
        return pd.DataFrame([])

    rows: list[dict[str, Any]] = []
    for doc in docs:
        patient_id = str(doc.get("_id") or "").strip()
        his_pid = str(doc.get("hisPid") or doc.get("hisPID") or "").strip()
        sex = str(doc.get("sex") or doc.get("gender") or "").strip().upper()
        if sex in {"男", "MALE"}:
            sex = "M"
        elif sex in {"女", "FEMALE"}:
            sex = "F"

        age = _as_number(doc.get("age"))
        if age is None:
            birthday = _parse_dt(doc.get("birthday"))
            if birthday:
                # birthday 已统一成带时区时间，当前时间也用 UTC，避免 naive/aware 相减报错。
                age = (datetime.now(timezone.utc) - birthday).days / 365.25
        
        if age is not None:
            age = int(math.floor(age))

        los_days = _infer_los_days(doc)
        icu_mortality = 1 if _infer_outcome(doc) == "dead" else 0
        hospital_mortality = _coerce_binary(doc.get("hospital_mortality"))
        mortality_28d = _coerce_binary(doc.get("mortality_28d") or doc.get("mortality28d"))

        row = {
            "patient_id": patient_id,
            "hisPid": his_pid,
            "name": doc.get("name"),
            "age": age,
            "sex": sex or None,
            "dept": doc.get("hisDept") or doc.get("dept"),
            "deptCode": doc.get("deptCode"),
            "outcome": _infer_outcome(doc),
            "discharge_dest": _infer_discharge_dest(doc),
            "primary_diagnosis": doc.get("clinicalDiagnosis") or doc.get("admissionDiagnosis") or doc.get("diagnosis") or "",
            "sofa_admission": _as_number(doc.get("sofa_admission")),
            "apache2": _as_number(doc.get("apache2")),
            "mechanical_ventilation": _coerce_binary(doc.get("mechanical_ventilation")),
            "crrt": _coerce_binary(doc.get("crrt")),
            "vasopressor": _coerce_binary(doc.get("vasopressor")),
            "los_icu_days": los_days,
            "los_icu_group": _infer_los_group(los_days),
            "icu_mortality": icu_mortality,
            "hospital_mortality": icu_mortality if hospital_mortality is None else hospital_mortality,
            "mortality_28d": icu_mortality if mortality_28d is None else mortality_28d,
            "admission_time": _parse_dt(doc.get("icuAdmissionTime") or doc.get("admissionTime")),
        }
        for key, value in doc.items():
            if key not in row:
                row[key] = value
        rows.append(row)

    df = pd.DataFrame(rows)
    # 统一补齐 SOFA/APACHE II，避免 patient 表缺失时分析模块大面积为空。
    try:
        await _attach_score_extrema(df, db, "sofa", "sofa_max")
        await _attach_score_extrema(df, db, "apacheII", "apache2_max")
        sofa_series = pd.to_numeric(df.get("sofa_admission"), errors="coerce")
        sofa_max_series = pd.to_numeric(df.get("sofa_max"), errors="coerce")
        apache_series = pd.to_numeric(df.get("apache2"), errors="coerce")
        apache_max_series = pd.to_numeric(df.get("apache2_max"), errors="coerce")
        df["sofa_admission"] = sofa_series.combine_first(sofa_max_series)
        df["apache2"] = apache_series.combine_first(apache_max_series)
        await _attach_score_extrema(df, db, ["gcs", "gcsScore"], "gcs_admission")
        await _attach_score_extrema(df, db, ["rass", "rassScore"], "rass_admission")
    except Exception as exc:
        logger.warning("attach score extrema failed: %s", exc)
    try:
        await _attach_lab_admission_values(df, db, LAB_NAME_MAP)
    except Exception as exc:
        logger.warning("attach lab admission values failed: %s", exc)
    try:
        await _attach_treatment_days(df, db)
    except Exception as exc:
        logger.warning("attach treatment days failed: %s", exc)
    return df


async def _resolve_patient_ids(
    patient_ids: list[str] | None,
    cohort_id: str | None,
    db,
) -> list[str]:
    resolved = _ensure_patient_id_set(patient_ids)
    if resolved:
        return resolved
    token = str(cohort_id or "").strip()
    if not token:
        return []

    query = {"cohort_id": token}
    oid = safe_oid(token)
    if oid is not None:
        query = {"$or": [{"_id": oid}, {"cohort_id": token}]}

    doc = await db.col("research_cohorts").find_one(query)
    if not doc:
        return []
    values: list[str] = []
    for key in ("patient_ids", "patients", "members"):
        items = doc.get(key)
        if isinstance(items, list):
            for item in items:
                text = str(item or "").strip()
                if text:
                    values.append(text)
    return _ensure_patient_id_set(values)


def _series_values(df: pd.DataFrame, field: str) -> pd.Series:
    if field not in df.columns:
        return pd.Series([], dtype=float)
    return df[field]


def _safe_float(value: Any) -> float | None:
    try:
        v = float(value)
    except Exception:
        return None
    return v if math.isfinite(v) else None


def _build_group_masks(
    df: pd.DataFrame,
    *,
    group_by: str | None,
    group_definitions: dict[str, Any] | None,
) -> list[tuple[str, pd.Series]]:
    if df.empty:
        return []
    if isinstance(group_definitions, dict) and group_definitions:
        rows: list[tuple[str, pd.Series]] = []
        for group_name, rule in group_definitions.items():
            label = str(group_name or "").strip() or "未命名分组"
            if isinstance(rule, dict):
                mask = _apply_filter(df, rule)
            else:
                mask = pd.Series([False] * len(df), index=df.index)
            rows.append((label, mask))
        if rows:
            return rows
    field = str(group_by or "").strip()
    if field and field in df.columns:
        groups: list[tuple[str, pd.Series]] = []
        values = sorted([str(v) for v in df[field].dropna().unique().tolist()])
        for value in values:
            groups.append((value, df[field].astype(str) == value))
        if groups:
            return groups
    return [("全部", pd.Series([True] * len(df), index=df.index))]


def _format_cat_cell(group_series: pd.Series, target: Any) -> str:
    total = int(group_series.notna().sum())
    if total <= 0:
        return "0 (0.0%)"
    count = int((group_series == target).sum())
    return _format_count_ratio(count, total)


def _format_multicat_cell(group_series: pd.Series, categories: list[Any]) -> str:
    total = int(group_series.notna().sum())
    if total <= 0:
        return "—"
    parts: list[str] = []
    for cat in categories:
        count = int((group_series == cat).sum())
        parts.append(f"{cat}:{count * 100.0 / total:.1f}%")
    return "; ".join(parts)


def _contingency_from_groups(groups: list[pd.Series], categories: list[Any]) -> list[list[int]]:
    table: list[list[int]] = []
    for cat in categories:
        row = [int((series == cat).sum()) for series in groups]
        table.append(row)
    return table


async def generate_table1(
    patient_ids: list[str],
    group_by: str,
    group_definitions: dict,
    variables: list[dict],
    db,
    *,
    cohort_id: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)

    if df.empty:
        return {
            "title": "Table 1. Baseline Characteristics",
            "groups": [],
            "rows": [],
            "footnote": "无可用数据",
            "n_total": 0,
        }

    group_rows = _build_group_masks(
        df,
        group_by=group_by,
        group_definitions=group_definitions,
    )
    group_labels: list[str] = []
    group_series: list[pd.DataFrame] = []
    for label, mask in group_rows:
        part = df[mask.fillna(False)]
        group_series.append(part)
        group_labels.append(f"{label} (n={len(part)})")

    rows: list[dict[str, Any]] = []
    for item in variables or []:
        field = str(item.get("field") or "").strip()
        if not field:
            continue
        label = str(item.get("label") or field)
        var_type = str(item.get("type") or "continuous").lower()
        source = _series_values(df, field)
        if source.empty:
            rows.append(
                {
                    "variable": label,
                    "type": var_type,
                    "values": ["—" for _ in group_series],
                    "statistic": "—",
                    "p_value": None,
                    "p_display": "—",
                    "significant": False,
                }
            )
            continue

        if var_type == "continuous":
            all_values = pd.to_numeric(source, errors="coerce").dropna().to_numpy(dtype=float)
            is_normal = _is_normal_distribution(all_values)
            values: list[str] = []
            arrays: list[Any] = []
            for group_df in group_series:
                arr = pd.to_numeric(group_df.get(field), errors="coerce").dropna().to_numpy(dtype=float)
                arrays.append(arr)
                values.append(_format_continuous(arr, is_normal))

            statistic = "—"
            p_value: float | None = None
            valid_arrays = [arr for arr in arrays if len(arr) > 0]
            if len(valid_arrays) >= 2:
                try:
                    if len(valid_arrays) == 2:
                        if is_normal:
                            t_value, p_value = stats.ttest_ind(valid_arrays[0], valid_arrays[1], equal_var=False)
                            statistic = f"t={t_value:.2f}"
                        else:
                            u_value, p_value = stats.mannwhitneyu(valid_arrays[0], valid_arrays[1], alternative="two-sided")
                            statistic = f"U={u_value:.1f}"
                    else:
                        if is_normal:
                            f_value, p_value = stats.f_oneway(*valid_arrays)
                            statistic = f"F={f_value:.2f}"
                        else:
                            h_value, p_value = stats.kruskal(*valid_arrays)
                            statistic = f"H={h_value:.2f}"
                except Exception:
                    p_value = None
                    statistic = "—"

            rows.append(
                {
                    "variable": label,
                    "field": field,
                    "type": "continuous",
                    "distribution": "normal" if is_normal else "non-normal",
                    "values": values,
                    "statistic": statistic,
                    "p_value": p_value,
                    "p_display": _p_display(p_value),
                    "significant": bool(p_value is not None and p_value < 0.05),
                }
            )
            continue

        series_groups = [group_df.get(field, pd.Series([], dtype=object)).astype(str) for group_df in group_series]
        raw_categories = item.get("categories")
        categories = raw_categories if isinstance(raw_categories, list) and raw_categories else sorted(
            [str(v) for v in source.dropna().astype(str).unique().tolist()]
        )
        if not categories:
            categories = ["1", "0"] if var_type == "binary" else ["是", "否"]

        p_value = None
        statistic = "—"
        table = _contingency_from_groups(series_groups, categories)
        if len(table) >= 2 and len(series_groups) >= 2:
            statistic, p_value = _chi_square_or_fisher(table)

        if var_type == "binary":
            target = item.get("positive")
            if target is None:
                target = categories[0] if categories else "1"
            values = [_format_cat_cell(series, str(target)) for series in series_groups]
        elif len(categories) <= 2:
            target = categories[0]
            values = [_format_cat_cell(series, target) for series in series_groups]
        else:
            show_categories = categories[:4]
            values = [_format_multicat_cell(series, show_categories) for series in series_groups]

        rows.append(
            {
                "variable": label,
                "field": field,
                "type": "binary" if var_type == "binary" else "categorical",
                "values": values,
                "statistic": statistic,
                "p_value": p_value,
                "p_display": _p_display(p_value),
                "significant": bool(p_value is not None and p_value < 0.05),
            }
        )

    return {
        "title": "Table 1. Baseline Characteristics",
        "groups": group_labels,
        "rows": rows,
        "n_total": int(len(df)),
        "group_by": group_by,
        "footnote": "连续变量正态分布用 mean±SD(t检验/ANOVA)；非正态用 median(IQR)(Mann-Whitney U/Kruskal-Wallis)；分类变量用 n(%) 表示(χ²/Fisher)。",
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


async def descriptive_statistics(
    patient_ids: list[str],
    variables: list[dict] | list[str],
    db,
    *,
    cohort_id: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)
    if df.empty:
        return {"n_total": 0, "variables": []}

    rows: list[dict[str, Any]] = []
    for item in variables or []:
        if isinstance(item, str):
            field = item
            label = item
            var_type = "auto"
        else:
            field = str(item.get("field") or "")
            label = str(item.get("label") or field)
            var_type = str(item.get("type") or "auto")
        if not field:
            continue
        series = df.get(field)
        if series is None:
            continue

        numeric = pd.to_numeric(series, errors="coerce").dropna()
        if var_type in {"continuous", "auto"} and len(numeric) >= max(3, int(len(series) * 0.3)):
            q1 = float(numeric.quantile(0.25))
            q2 = float(numeric.quantile(0.5))
            q3 = float(numeric.quantile(0.75))
            rows.append(
                {
                    "field": field,
                    "label": label,
                    "type": "continuous",
                    "n": int(numeric.count()),
                    "mean": float(numeric.mean()),
                    "std": float(numeric.std(ddof=1)) if numeric.count() > 1 else 0.0,
                    "median": q2,
                    "q1": q1,
                    "q3": q3,
                    "min": float(numeric.min()),
                    "max": float(numeric.max()),
                    "normal": _is_normal_distribution(numeric.to_numpy(dtype=float)),
                }
            )
            continue

        value_counts = series.astype(str).value_counts(dropna=True)
        top = value_counts.head(10)
        rows.append(
            {
                "field": field,
                "label": label,
                "type": "categorical",
                "n": int(series.notna().sum()),
                "categories": [
                    {
                        "name": str(k),
                        "count": int(v),
                        "ratio": float(v / max(1, series.notna().sum())),
                    }
                    for k, v in top.items()
                ],
            }
        )

    return {
        "n_total": int(len(df)),
        "variables": rows,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


def _coerce_series_binary(series: pd.Series) -> pd.Series:
    return series.apply(_coerce_binary).dropna().astype(int)


def _coerce_series_numeric(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce").dropna().astype(float)


def _km_curve_fallback(part: pd.DataFrame) -> tuple[dict[str, Any], float | None]:
    if part.empty:
        return (
            {
                "timeline": [],
                "survival": [],
                "ci_lower": [],
                "ci_upper": [],
                "at_risk": [],
                "n": 0,
                "events": 0,
            },
            None,
        )
    work = part[["_time", "_event"]].copy()
    work["_time"] = pd.to_numeric(work["_time"], errors="coerce")
    work["_event"] = work["_event"].apply(_coerce_binary)
    work = work.dropna(subset=["_time", "_event"]).sort_values("_time")
    if work.empty:
        return (
            {
                "timeline": [],
                "survival": [],
                "ci_lower": [],
                "ci_upper": [],
                "at_risk": [],
                "n": 0,
                "events": 0,
            },
            None,
        )
    event_table = (
        work.groupby("_time")
        .agg(events=("_event", "sum"), total=("_event", "count"))
        .reset_index()
        .sort_values("_time")
    )
    event_table["censored"] = event_table["total"] - event_table["events"]
    n_at_risk = int(len(work))
    surv = 1.0
    greenwood = 0.0
    timeline = [0.0]
    survival_values = [1.0]
    ci_lower = [1.0]
    ci_upper = [1.0]
    at_risk = [n_at_risk]
    median_survival: float | None = None
    for _, row in event_table.iterrows():
        t = float(row["_time"])
        d = float(row["events"])
        c = float(row["censored"])
        if n_at_risk <= 0:
            break
        if d > 0:
            surv *= max(0.0, 1.0 - (d / n_at_risk))
            if n_at_risk - d > 0:
                greenwood += d / (n_at_risk * (n_at_risk - d))
            se = math.sqrt(max(0.0, (surv ** 2) * greenwood))
            lo = max(0.0, surv - 1.96 * se)
            hi = min(1.0, surv + 1.96 * se)
            timeline.append(t)
            survival_values.append(float(surv))
            ci_lower.append(float(lo))
            ci_upper.append(float(hi))
            at_risk.append(int(n_at_risk))
            if median_survival is None and surv <= 0.5:
                median_survival = t
        n_at_risk -= int(d + c)
    return (
        {
            "timeline": timeline,
            "survival": survival_values,
            "ci_lower": ci_lower,
            "ci_upper": ci_upper,
            "at_risk": at_risk,
            "n": int(len(work)),
            "events": int(work["_event"].sum()),
        },
        median_survival,
    )


async def survival_analysis(
    patient_ids: list[str],
    time_field: str,
    event_field: str,
    group_by: str | None = None,
    max_time: int = 28,
    db=None,
    *,
    cohort_id: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)
    if df.empty:
        return {"kaplan_meier": {"curves": {}, "log_rank_p": None, "median_survival": {}}, "cox_regression": None}

    if time_field not in df.columns:
        df[time_field] = pd.to_numeric(df.get("los_icu_days"), errors="coerce")
    if event_field not in df.columns:
        df[event_field] = df.get("icu_mortality", 0)

    df = df.copy()
    df["_time"] = pd.to_numeric(df.get(time_field), errors="coerce")
    df["_event"] = df.get(event_field).apply(_coerce_binary)
    df = df.dropna(subset=["_time", "_event"])
    df = df[(df["_time"] >= 0) & (df["_time"] <= float(max_time))]
    if df.empty:
        return {
            "kaplan_meier": {"curves": {}, "log_rank_p": None, "median_survival": {}},
            "cox_regression": None,
            "n_total": 0,
            "n_events": 0,
            "reason": "no_valid_time_event_data",
        }

    if group_by and group_by in df.columns:
        group_values = sorted([str(v) for v in df[group_by].dropna().unique().tolist()])
    else:
        group_values = ["全部"]
        df["_all_group"] = "全部"
        group_by = "_all_group"

    curves: dict[str, Any] = {}
    median_survival: dict[str, float | None] = {}
    for group in group_values:
        part = df[df[group_by].astype(str) == str(group)]
        if part.empty:
            continue
        if KaplanMeierFitter is not None:
            kmf = KaplanMeierFitter()
            kmf.fit(part["_time"], event_observed=part["_event"], label=group)
            sf = kmf.survival_function_
            ci = kmf.confidence_interval_
            timeline = [float(x) for x in sf.index.tolist()]
            survival_values = [float(v) for v in sf.iloc[:, 0].tolist()]
            ci_lower = [float(v) for v in ci.iloc[:, 0].tolist()]
            ci_upper = [float(v) for v in ci.iloc[:, 1].tolist()]
            risk = kmf.event_table.get("at_risk")
            at_risk = [int(v) for v in risk.tolist()] if risk is not None else [0 for _ in timeline]
            median = kmf.median_survival_time_
            median_survival[group] = None if median is None or not math.isfinite(float(median)) else float(median)
            curves[group] = {
                "timeline": timeline,
                "survival": survival_values,
                "ci_lower": ci_lower,
                "ci_upper": ci_upper,
                "at_risk": at_risk,
                "n": int(len(part)),
                "events": int(part["_event"].sum()),
            }
        else:
            curve, median = _km_curve_fallback(part)
            curves[group] = curve
            median_survival[group] = median

    log_rank_p = None
    if KaplanMeierFitter is not None and len(group_values) >= 2 and logrank_test is not None:
        try:
            if len(group_values) == 2:
                g1 = df[df[group_by].astype(str) == group_values[0]]
                g2 = df[df[group_by].astype(str) == group_values[1]]
                lr = logrank_test(g1["_time"], g2["_time"], g1["_event"], g2["_event"])
                log_rank_p = float(lr.p_value)
            elif multivariate_logrank_test is not None:
                lr = multivariate_logrank_test(df["_time"], df[group_by].astype(str), df["_event"])
                log_rank_p = float(lr.p_value)
        except Exception:
            log_rank_p = None

    cox_result: dict[str, Any] | None = None
    if KaplanMeierFitter is not None and len(group_values) >= 2 and CoxPHFitter is not None:
        try:
            cox_df = df[["_time", "_event", group_by]].copy()
            dummies = pd.get_dummies(cox_df[group_by].astype(str), prefix="group", drop_first=True)
            cox_input = pd.concat([cox_df[["_time", "_event"]], dummies], axis=1).dropna()
            if len(cox_input) >= 20 and len(dummies.columns) >= 1:
                cph = CoxPHFitter()
                cph.fit(cox_input, duration_col="_time", event_col="_event")
                summary = cph.summary.reset_index()
                item = summary.iloc[0]
                hr = float(item.get("exp(coef)", item.get("exp_coef", math.nan)))
                ci_low = float(item.get("exp(coef) lower 95%", item.get("exp_coef_lower_95", math.nan)))
                ci_up = float(item.get("exp(coef) upper 95%", item.get("exp_coef_upper_95", math.nan)))
                p_value = float(item.get("p", math.nan))
                cox_result = {
                    "variable": str(item.get("covariate", item.get("index", "group"))),
                    "hr": hr,
                    "ci_lower": ci_low,
                    "ci_upper": ci_up,
                    "p_value": p_value if math.isfinite(p_value) else None,
                    "p_display": _p_display(p_value if math.isfinite(p_value) else None),
                    "c_index": float(getattr(cph, "concordance_index_", 0.0)),
                }
        except Exception:
            cox_result = None

    return {
        "kaplan_meier": {
            "curves": curves,
            "log_rank_p": log_rank_p,
            "log_rank_p_display": _p_display(log_rank_p),
            "median_survival": median_survival,
        },
        "cox_regression": cox_result,
        "n_total": int(len(df)),
        "n_events": int(df["_event"].sum()),
        "dependency_mode": "lifelines" if KaplanMeierFitter is not None else "fallback",
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


def _encode_numeric_predictor(series: pd.Series) -> pd.Series:
    numeric = pd.to_numeric(series, errors="coerce")
    if numeric.notna().sum() >= max(3, int(0.3 * len(series))):
        return numeric
    cat = series.where(series.notna(), pd.NA)
    cat = cat.astype(str).str.strip().str.lower()
    cat = cat.replace({"": pd.NA, "nan": pd.NA, "none": pd.NA, "null": pd.NA, "nat": pd.NA, "<na>": pd.NA})
    codes, _ = pd.factorize(cat)
    out = pd.Series(codes, index=series.index, dtype=float)
    out[out < 0] = math.nan
    return out


def _logistic_hosmer_lemeshow(y_true: pd.Series, y_prob: pd.Series, g: int = 10) -> float | None:
    try:
        df = pd.DataFrame({"y": y_true, "p": y_prob}).dropna()
        if len(df) < 20:
            return None
        df["bucket"] = pd.qcut(df["p"], q=min(g, len(df) // 2), duplicates="drop")
        grouped = df.groupby("bucket")
        obs = grouped["y"].sum()
        exp = grouped["p"].sum()
        n = grouped["y"].count()
        stat = (((obs - exp) ** 2) / (exp + 1e-9) + (((n - obs) - (n - exp)) ** 2) / ((n - exp) + 1e-9)).sum()
        dof = max(1, len(obs) - 2)
        p_value = 1 - stats.chi2.cdf(stat, dof)
        return float(p_value)
    except Exception:
        return None


def _fit_binary_logistic_sklearn(work: pd.DataFrame, predictors: list[str]) -> tuple[list[dict[str, Any]], Any]:
    if LogisticRegression is None:
        raise RuntimeError("缺少 sklearn 依赖，无法回退执行 Logistic 回归")
    if work.empty or work.iloc[:, 0].nunique() < 2:
        return [], None
    y = work.iloc[:, 0].astype(int)
    x = work[predictors].astype(float)
    model = LogisticRegression(max_iter=1000)
    model.fit(x, y)
    probs = model.predict_proba(x)[:, 1]
    rows: list[dict[str, Any]] = []
    for idx, col in enumerate(predictors):
        beta = float(model.coef_[0][idx])
        rows.append(
            {
                "variable": col,
                "estimate": float(math.exp(beta)),
                "ci_lower": None,
                "ci_upper": None,
                "p": None,
                "p_display": "—",
                "label": "OR",
            }
        )
    return rows, {"engine": "sklearn_logistic", "metrics": {"auc": float(metrics.roc_auc_score(y, probs)) if len(set(y.tolist())) >= 2 else None}}


def _fit_binary_logistic(df: pd.DataFrame, outcome_col: str, predictors: list[str]) -> tuple[list[dict[str, Any]], Any]:
    work = df[[outcome_col, *predictors]].copy()
    work[outcome_col] = work[outcome_col].apply(_coerce_binary)
    for col in predictors:
        work[col] = _encode_numeric_predictor(work[col])
    work = work.dropna()
    if work.empty or work[outcome_col].nunique() < 2:
        return [], None
    if sm is None:
        if LogisticRegression is None:
            raise RuntimeError("缺少 statsmodels/sklearn 依赖，无法执行 Logistic 回归")
        return _fit_binary_logistic_sklearn(work, predictors)
    y = work[outcome_col].astype(int)
    x = sm.add_constant(work[predictors], has_constant="add")
    try:
        model = sm.Logit(y, x).fit(disp=False, maxiter=200)
        conf = model.conf_int()
        rows: list[dict[str, Any]] = []
        for col in predictors:
            beta = float(model.params[col])
            p = float(model.pvalues[col])
            rows.append(
                {
                    "variable": col,
                    "estimate": float(math.exp(beta)),
                    "ci_lower": float(math.exp(conf.loc[col, 0])),
                    "ci_upper": float(math.exp(conf.loc[col, 1])),
                    "p": p,
                    "p_display": _p_display(p),
                    "label": "OR",
                }
            )
        return rows, model
    except Exception:
        return _fit_binary_logistic_sklearn(work, predictors)


def _fit_linear(df: pd.DataFrame, outcome_col: str, predictors: list[str]) -> tuple[list[dict[str, Any]], Any]:
    if sm is None:
        if LinearRegression is None:
            raise RuntimeError("缺少 statsmodels/sklearn 依赖，无法执行线性回归")
        work = df[[outcome_col, *predictors]].copy()
        work[outcome_col] = pd.to_numeric(work[outcome_col], errors="coerce")
        for col in predictors:
            work[col] = _encode_numeric_predictor(work[col])
        work = work.dropna()
        if work.empty:
            return [], None
        y = work[outcome_col].astype(float)
        x = work[predictors].astype(float)
        model = LinearRegression()
        model.fit(x, y)
        pred = model.predict(x)
        ss_res = float(((y - pred) ** 2).sum())
        ss_tot = float(((y - float(y.mean())) ** 2).sum())
        r2 = None if ss_tot <= 0 else float(1 - ss_res / ss_tot)
        rows: list[dict[str, Any]] = []
        for idx, col in enumerate(predictors):
            rows.append(
                {
                    "variable": col,
                    "estimate": float(model.coef_[idx]),
                    "ci_lower": None,
                    "ci_upper": None,
                    "p": None,
                    "p_display": "—",
                    "label": "β",
                }
            )
        return rows, {"engine": "sklearn_linear", "metrics": {"r2": r2}}
    work = df[[outcome_col, *predictors]].copy()
    work[outcome_col] = pd.to_numeric(work[outcome_col], errors="coerce")
    for col in predictors:
        work[col] = _encode_numeric_predictor(work[col])
    work = work.dropna()
    if work.empty:
        return [], None
    y = work[outcome_col].astype(float)
    x = sm.add_constant(work[predictors], has_constant="add")
    model = sm.OLS(y, x).fit()
    conf = model.conf_int()
    rows: list[dict[str, Any]] = []
    for col in predictors:
        beta = float(model.params[col])
        p = float(model.pvalues[col])
        rows.append(
            {
                "variable": col,
                "estimate": beta,
                "ci_lower": float(conf.loc[col, 0]),
                "ci_upper": float(conf.loc[col, 1]),
                "p": p,
                "p_display": _p_display(p),
                "label": "β",
            }
        )
    return rows, model


def _fit_cox(df: pd.DataFrame, time_col: str, event_col: str, predictors: list[str]) -> tuple[list[dict[str, Any]], Any]:
    if CoxPHFitter is None:
        raise RuntimeError("缺少 lifelines 依赖，无法执行 Cox 回归")
    work = df[[time_col, event_col, *predictors]].copy()
    work[time_col] = pd.to_numeric(work[time_col], errors="coerce")
    work[event_col] = work[event_col].apply(_coerce_binary)
    for col in predictors:
        work[col] = _encode_numeric_predictor(work[col])
    work = work.dropna()
    work = work[work[time_col] >= 0]
    if work.empty:
        return [], None
    cph = CoxPHFitter()
    cph.fit(work, duration_col=time_col, event_col=event_col)
    summary = cph.summary
    rows: list[dict[str, Any]] = []
    for col in predictors:
        if col not in summary.index:
            continue
        row = summary.loc[col]
        p = float(row.get("p", math.nan))
        rows.append(
            {
                "variable": col,
                "estimate": float(row.get("exp(coef)", math.nan)),
                "ci_lower": float(row.get("exp(coef) lower 95%", math.nan)),
                "ci_upper": float(row.get("exp(coef) upper 95%", math.nan)),
                "p": p if math.isfinite(p) else None,
                "p_display": _p_display(p if math.isfinite(p) else None),
                "label": "HR",
            }
        )
    return rows, cph


async def regression_analysis(
    patient_ids: list[str],
    outcome: str,
    outcome_type: str,
    predictors: list[str],
    time_field: str | None = None,
    confounders: list[str] | None = None,
    db=None,
    *,
    cohort_id: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)
    if df.empty:
        return {"model_type": outcome_type, "univariate": [], "multivariate": [], "model_metrics": {}}

    outcome_key = str(outcome or "").strip()
    if outcome_key not in df.columns:
        if outcome_type == "binary":
            df[outcome_key] = df.get("icu_mortality", 0)
        elif outcome_type == "continuous":
            df[outcome_key] = pd.to_numeric(df.get("los_icu_days"), errors="coerce")
        else:
            df[outcome_key] = df.get("icu_mortality", 0)

    clean_predictors = [str(x).strip() for x in predictors or [] if str(x).strip() and str(x).strip() in df.columns]
    if not clean_predictors:
        return {"model_type": outcome_type, "univariate": [], "multivariate": [], "model_metrics": {}, "reason": "no_valid_predictors", "n_total": int(len(df))}

    model_type = str(outcome_type or "binary").lower()
    outcome_binary = df[outcome_key].apply(_coerce_binary) if outcome_key in df.columns else pd.Series([], dtype=float)
    outcome_non_null = int(outcome_binary.dropna().shape[0]) if model_type == "binary" else int(pd.to_numeric(df[outcome_key], errors="coerce").dropna().shape[0])
    outcome_unique = int(outcome_binary.dropna().nunique()) if model_type == "binary" else None
    if model_type == "binary" and (outcome_unique or 0) < 2:
        return {
            "model_type": "logistic",
            "univariate": [],
            "multivariate": [],
            "model_metrics": {},
            "n_total": int(len(df)),
            "reason": "outcome_single_class",
            "outcome_non_null": outcome_non_null,
            "outcome_positive": int((outcome_binary == 1).sum()),
        }
    univariate_rows: list[dict[str, Any]] = []
    selected_for_multi: list[str] = []
    univariate_counts: list[dict[str, Any]] = []

    for pred in clean_predictors:
        if model_type == "binary":
            count_df = pd.DataFrame({
                "_outcome": df[outcome_key].apply(_coerce_binary),
                "_predictor": _encode_numeric_predictor(df[pred]),
            }).dropna()
        else:
            count_df = pd.DataFrame({
                "_outcome": pd.to_numeric(df[outcome_key], errors="coerce"),
                "_predictor": _encode_numeric_predictor(df[pred]),
            }).dropna()
        univariate_counts.append({
            "variable": pred,
            "n_model": int(len(count_df)),
            "n_excluded": int(len(df) - len(count_df)),
        })
        try:
            if model_type == "binary":
                rows, _ = _fit_binary_logistic(df, outcome_key, [pred])
            elif model_type == "continuous":
                rows, _ = _fit_linear(df, outcome_key, [pred])
            else:
                if not time_field:
                    raise RuntimeError("survival 回归缺少 time_field")
                if time_field not in df.columns:
                    df[time_field] = pd.to_numeric(df.get("los_icu_days"), errors="coerce")
                rows, _ = _fit_cox(df, time_field, outcome_key, [pred])
            if not rows:
                continue
            row = rows[0]
            univariate_rows.append(row)
            p_value = _safe_float(row.get("p"))
            if p_value is not None and p_value < 0.1:
                selected_for_multi.append(pred)
        except Exception:
            continue

    final_predictors = list(dict.fromkeys([*selected_for_multi, *(confounders or [])]))
    final_predictors = [x for x in final_predictors if x in df.columns]
    multivariate_rows: list[dict[str, Any]] = []
    model_metrics: dict[str, Any] = {}
    multivariate_count: dict[str, Any] = {"n_model": 0, "n_excluded": int(len(df)), "variables": final_predictors}

    if final_predictors:
        if model_type == "binary":
            count_df = pd.DataFrame({"_outcome": df[outcome_key].apply(_coerce_binary)})
        else:
            count_df = pd.DataFrame({"_outcome": pd.to_numeric(df[outcome_key], errors="coerce")})
        for pred in final_predictors:
            count_df[pred] = _encode_numeric_predictor(df[pred])
        count_df = count_df.dropna()
        multivariate_count = {
            "n_model": int(len(count_df)),
            "n_excluded": int(len(df) - len(count_df)),
            "variables": final_predictors,
        }
        try:
            if model_type == "binary":
                multivariate_rows, model = _fit_binary_logistic(df, outcome_key, final_predictors)
                if model is not None:
                    if isinstance(model, dict):
                        model_metrics = model.get("metrics", {}) if isinstance(model.get("metrics"), dict) else {}
                    else:
                        pred = model.predict()
                        hl_p = _logistic_hosmer_lemeshow(model.model.endog, pred)
                        model_metrics = {
                            "aic": float(model.aic),
                            "bic": float(model.bic),
                            "pseudo_r2": float(getattr(model, "prsquared", math.nan)),
                            "hosmer_lemeshow_p": hl_p,
                        }
                    for row in multivariate_rows:
                        row["label"] = "aOR"
            elif model_type == "continuous":
                multivariate_rows, model = _fit_linear(df, outcome_key, final_predictors)
                if model is not None:
                    if isinstance(model, dict):
                        model_metrics = model.get("metrics", {}) if isinstance(model.get("metrics"), dict) else {}
                    else:
                        model_metrics = {
                            "aic": float(model.aic),
                            "bic": float(model.bic),
                            "r2": float(model.rsquared),
                            "adj_r2": float(model.rsquared_adj),
                        }
            else:
                if not time_field:
                    raise RuntimeError("survival 回归缺少 time_field")
                multivariate_rows, model = _fit_cox(df, time_field, outcome_key, final_predictors)
                if model is not None:
                    model_metrics = {
                        "aic": float(model.AIC_partial_),
                        "c_index": float(model.concordance_index_),
                    }
                    for row in multivariate_rows:
                        row["label"] = "aHR"
        except Exception:
            multivariate_rows = []
            model_metrics = {}

    for row in univariate_rows:
        row["p_display"] = _p_display(_safe_float(row.get("p")))
    for row in multivariate_rows:
        row["p_display"] = _p_display(_safe_float(row.get("p")))

    return {
        "model_type": {"binary": "logistic", "continuous": "linear", "survival": "cox"}.get(model_type, model_type),
        "univariate": univariate_rows,
        "multivariate": multivariate_rows,
        "model_metrics": model_metrics,
        "n_total": int(len(df)),
        "outcome_non_null": outcome_non_null,
        "univariate_counts": univariate_counts,
        "multivariate_count": multivariate_count,
        "reason": "ok" if (univariate_rows or multivariate_rows) else "no_model_fit",
        "outcome_positive": int((outcome_binary == 1).sum()) if model_type == "binary" else None,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


def _bootstrap_auc_ci(y_true: Any, scores: Any, *, n_boot: int = 300, alpha: float = 0.95) -> tuple[float | None, float | None]:
    if np is None:
        return None, None
    y = np.asarray(y_true, dtype=int)
    s = np.asarray(scores, dtype=float)
    if len(y) < 10:
        return None, None
    rng = np.random.default_rng(42)
    aucs: list[float] = []
    for _ in range(n_boot):
        idx = rng.integers(0, len(y), len(y))
        yb = y[idx]
        sb = s[idx]
        if len(np.unique(yb)) < 2:
            continue
        aucs.append(float(metrics.roc_auc_score(yb, sb)))
    if not aucs:
        return None, None
    lower_q = (1 - alpha) / 2
    upper_q = 1 - lower_q
    return float(np.quantile(aucs, lower_q)), float(np.quantile(aucs, upper_q))


def _bootstrap_auc_diff_test(y_true: Any, score_a: Any, score_b: Any, *, n_boot: int = 500) -> tuple[float | None, float | None]:
    if np is None:
        return None, None
    y = np.asarray(y_true, dtype=int)
    a = np.asarray(score_a, dtype=float)
    b = np.asarray(score_b, dtype=float)
    if len(y) < 20:
        return None, None
    rng = np.random.default_rng(7)
    diffs: list[float] = []
    for _ in range(n_boot):
        idx = rng.integers(0, len(y), len(y))
        yb = y[idx]
        if len(np.unique(yb)) < 2:
            continue
        dif = float(metrics.roc_auc_score(yb, a[idx]) - metrics.roc_auc_score(yb, b[idx]))
        diffs.append(dif)
    if len(diffs) < 20:
        return None, None
    mean = float(np.mean(diffs))
    std = float(np.std(diffs, ddof=1))
    if std <= 0:
        return None, None
    z_value = mean / std
    p_value = float(2 * (1 - stats.norm.cdf(abs(z_value))))
    return z_value, p_value


async def roc_analysis(
    patient_ids: list[str],
    outcome: str,
    predictors: list[str],
    db=None,
    *,
    cohort_id: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)
    if df.empty:
        return {"curves": {}, "delong_test": {}, "n_total": 0, "reason": "empty_cohort"}

    outcome_key = str(outcome or "").strip()
    if outcome_key not in df.columns:
        df[outcome_key] = df.get("icu_mortality", 0)
    y_series = df[outcome_key].apply(_coerce_binary).dropna().astype(int)
    if y_series.empty or y_series.nunique() < 2:
        return {
            "curves": {},
            "delong_test": {},
            "n_total": int(len(df)),
            "outcome_positive": int((y_series == 1).sum()) if not y_series.empty else 0,
            "reason": "outcome_single_class",
        }

    curves: dict[str, Any] = {}
    usable_scores: dict[str, Any] = {}

    for raw_pred in predictors or []:
        pred = str(raw_pred or "").strip()
        if not pred or pred not in df.columns:
            continue
        joined = pd.DataFrame({
            "y": df[outcome_key].apply(_coerce_binary),
            "score": pd.to_numeric(df[pred], errors="coerce"),
        }).dropna()
        if len(joined) < 10 or joined["y"].nunique() < 2:
            continue
        y = joined["y"].astype(int).to_numpy()
        score = joined["score"].astype(float).to_numpy()
        fpr, tpr, thresholds = metrics.roc_curve(y, score)
        auc = float(metrics.auc(fpr, tpr))
        ci_lower, ci_upper = _bootstrap_auc_ci(y, score)
        youden = tpr - fpr
        best_idx = int(np.argmax(youden)) if np is not None and len(youden) else 0
        cutoff = float(thresholds[best_idx]) if len(thresholds) > best_idx else None
        curves[pred] = {
            "fpr": [float(x) for x in fpr.tolist()],
            "tpr": [float(x) for x in tpr.tolist()],
            "auc": auc,
            "ci_lower": ci_lower,
            "ci_upper": ci_upper,
            "optimal_cutoff": cutoff,
            "sensitivity_at_cutoff": float(tpr[best_idx]) if len(tpr) > best_idx else None,
            "specificity_at_cutoff": float(1 - fpr[best_idx]) if len(fpr) > best_idx else None,
            "n": int(len(joined)),
        }
        usable_scores[pred] = (y, score)

    delong_test: dict[str, Any] = {}
    names = list(usable_scores.keys())
    for idx_a, name_a in enumerate(names):
        for name_b in names[idx_a + 1 :]:
            ya, sa = usable_scores[name_a]
            yb, sb = usable_scores[name_b]
            min_len = min(len(ya), len(yb))
            if min_len < 20:
                continue
            z_value, p_value = _bootstrap_auc_diff_test(ya[:min_len], sa[:min_len], sb[:min_len])
            delong_test[f"{name_a} vs {name_b}"] = {
                "z": z_value,
                "p": p_value,
                "p_display": _p_display(p_value),
            }

    return {
        "curves": curves,
        "delong_test": delong_test,
        "n_total": int(len(df)),
        "outcome_positive": int((y_series == 1).sum()),
        "reason": "ok" if curves else "no_valid_predictor_pairs",
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


async def subgroup_analysis(
    patient_ids: list[str],
    exposure: str,
    outcome: str,
    outcome_type: str,
    subgroups: list[dict],
    db=None,
    *,
    cohort_id: str | None = None,
    time_field: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)
    if df.empty or exposure not in df.columns:
        return {"results": [], "interaction_p": {}}

    results: list[dict[str, Any]] = []
    for item in subgroups or []:
        name = str(item.get("name") or "亚组")
        rule = item.get("filter") if isinstance(item.get("filter"), dict) else {}
        part = df[_apply_filter(df, rule)]
        if part.empty:
            results.append(
                {
                    "subgroup": name,
                    "n": 0,
                    "n_event": 0,
                    "estimate": None,
                    "ci_lower": None,
                    "ci_upper": None,
                    "p": None,
                    "p_display": "—",
                }
            )
            continue
        try:
            if outcome_type == "survival":
                if not time_field:
                    raise RuntimeError("survival 亚组分析缺少 time_field")
                rows, _ = _fit_cox(part, time_field, outcome, [exposure])
            else:
                rows, _ = _fit_binary_logistic(part, outcome, [exposure])
            row = rows[0] if rows else {}
            p_value = _safe_float(row.get("p"))
            n_event = int(part[outcome].apply(_coerce_binary).fillna(0).sum()) if outcome in part.columns else 0
            results.append(
                {
                    "subgroup": name,
                    "n": int(len(part)),
                    "n_event": n_event,
                    "estimate": _safe_float(row.get("estimate")),
                    "ci_lower": _safe_float(row.get("ci_lower")),
                    "ci_upper": _safe_float(row.get("ci_upper")),
                    "p": p_value,
                    "p_display": _p_display(p_value),
                }
            )
        except Exception:
            results.append(
                {
                    "subgroup": name,
                    "n": int(len(part)),
                    "n_event": int(part[outcome].apply(_coerce_binary).fillna(0).sum()) if outcome in part.columns else 0,
                    "estimate": None,
                    "ci_lower": None,
                    "ci_upper": None,
                    "p": None,
                    "p_display": "—",
                }
            )

    interaction_p: dict[str, Any] = {}
    if outcome_type != "survival" and sm is not None and len(subgroups or []) >= 2 and outcome in df.columns:
        for item in (subgroups or []):
            name = str(item.get("name") or "").strip()
            rule = item.get("filter") if isinstance(item.get("filter"), dict) else {}
            if not name:
                continue
            mask = _apply_filter(df, rule).fillna(False)
            if mask.nunique() <= 1:
                continue
            try:
                tmp = df[[outcome, exposure]].copy()
                tmp["_sub"] = mask.astype(int)
                tmp[outcome] = tmp[outcome].apply(_coerce_binary)
                tmp[exposure] = _encode_numeric_predictor(tmp[exposure])
                tmp = tmp.dropna()
                if tmp.empty:
                    continue
                tmp["_interaction"] = tmp[exposure] * tmp["_sub"]
                y = tmp[outcome].astype(int)
                x = sm.add_constant(tmp[[exposure, "_sub", "_interaction"]], has_constant="add")
                model = sm.Logit(y, x).fit(disp=False, maxiter=200)
                p_value = _safe_float(model.pvalues.get("_interaction"))
                interaction_p[name] = p_value
            except Exception:
                interaction_p[name] = None

    return {
        "results": results,
        "subgroups": results,
        "interaction_p": interaction_p,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


async def correlation_analysis(
    patient_ids: list[str],
    variables: list[str],
    method: str = "auto",
    db=None,
    *,
    cohort_id: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)
    if df.empty:
        return {
            "method": method,
            "matrix": {"labels": [], "correlations": [], "p_values": [], "n_pairs": []},
            "variable_coverage": [],
            "excluded_variables": [],
        }

    requested = [str(v).strip() for v in variables or [] if str(v).strip()]
    labels = [field for field in requested if field in df.columns]
    coverage_rows: list[dict[str, Any]] = []
    excluded_variables: list[dict[str, Any]] = []
    total_count = int(len(df))
    for field in labels:
        series = pd.to_numeric(df[field], errors="coerce")
        non_null_count = int(series.notna().sum())
        coverage_rows.append(
            {
                "field": field,
                "non_null_count": non_null_count,
                "total_count": total_count,
                "non_null_rate": (non_null_count / total_count) if total_count else None,
            }
        )
        if non_null_count < 3:
            excluded_variables.append(
                {
                    "field": field,
                    "reason": "insufficient_non_null",
                    "non_null_count": non_null_count,
                    "total_count": total_count,
                }
            )
    labels = [field for field in labels if field not in {item["field"] for item in excluded_variables}]
    if len(labels) < 2:
        return {
            "method": method,
            "matrix": {"labels": labels, "correlations": [], "p_values": [], "n_pairs": []},
            "variable_coverage": coverage_rows,
            "excluded_variables": excluded_variables,
        }

    numeric_df = pd.DataFrame({label: pd.to_numeric(df[label], errors="coerce") for label in labels})
    chosen = method.lower()
    if chosen == "auto":
        chosen = "pearson"
        for label in labels:
            series = numeric_df[label].dropna().to_numpy(dtype=float)
            if len(series) >= 3 and not _is_normal_distribution(series):
                chosen = "spearman"
                break

    def _compute_sync() -> tuple[list[list[float]], list[list[float | None]], list[list[int]]]:
        corr_matrix: list[list[float]] = []
        p_matrix: list[list[float | None]] = []
        n_matrix: list[list[int]] = []
        for i, label_i in enumerate(labels):
            corr_row: list[float] = []
            p_row: list[float | None] = []
            n_row: list[int] = []
            for j, label_j in enumerate(labels):
                if i == j:
                    corr_row.append(1.0)
                    p_row.append(None)
                    n_row.append(int(numeric_df[label_i].dropna().shape[0]))
                    continue
                pair = numeric_df[[label_i, label_j]].dropna()
                n_row.append(int(len(pair)))
                if len(pair) < 3:
                    corr_row.append(0.0)
                    p_row.append(None)
                    continue
                try:
                    if chosen == "pearson":
                        corr, p_val = stats.pearsonr(pair[label_i], pair[label_j])
                    else:
                        corr, p_val = stats.spearmanr(pair[label_i], pair[label_j])
                    corr_row.append(float(corr))
                    p_row.append(float(p_val))
                except Exception:
                    corr_row.append(0.0)
                    p_row.append(None)
            corr_matrix.append(corr_row)
            p_matrix.append(p_row)
            n_matrix.append(n_row)
        return corr_matrix, p_matrix, n_matrix

    corr_matrix, p_matrix, n_matrix = await asyncio.to_thread(_compute_sync)

    return {
        "method": chosen,
        "matrix": {
            "labels": labels,
            "correlations": corr_matrix,
            "p_values": p_matrix,
            "n_pairs": n_matrix,
        },
        "variable_coverage": coverage_rows,
        "excluded_variables": excluded_variables,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


def _indicator_codes(indicator: str) -> list[str]:
    key = str(indicator or "").strip().lower()
    return VITAL_CODE_MAP.get(key, [indicator])


def _extract_bedside_value(doc: dict[str, Any]) -> float | None:
    for key in ("fVal", "intVal", "strVal", "value"):
        value = _as_number(doc.get(key))
        if value is not None:
            return value
    return None


def _extract_lab_value(doc: dict[str, Any]) -> float | None:
    for key in ("result", "itemValue", "value"):
        value = _as_number(doc.get(key))
        if value is not None:
            return value
    return None


def _extract_lab_time(doc: dict[str, Any]) -> datetime | None:
    return _parse_dt(doc.get("authTime") or doc.get("collectTime") or doc.get("testTime") or doc.get("reportTime") or doc.get("time"))


def _trend_indicator_lab_key(indicator: str) -> str | None:
    key = str(indicator or "").strip().lower()
    return TREND_LAB_FIELD_MAP.get(key)


async def trend_analysis(
    patient_ids: list[str],
    indicators: list[str],
    time_reference: str,
    time_range_hours: int,
    group_by: str | None = None,
    interval_hours: float = 4,
    db=None,
    *,
    cohort_id: str | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved_ids = await _resolve_patient_ids(patient_ids, cohort_id, db)
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved_ids, db, max_patients=max_patients)
    if df.empty:
        return {"indicators": {}}

    group_col = str(group_by or "").strip()
    if not group_col or group_col not in df.columns:
        df["_group"] = "全部"
        group_col = "_group"
    groups = sorted([str(v) for v in df[group_col].fillna("未分组").astype(str).unique().tolist()])

    interval = max(0.5, float(interval_hours or 4))
    max_hours = max(1.0, float(time_range_hours or 72))
    timeline = [round(v, 3) for v in np.arange(0, max_hours + 1e-9, interval).tolist()] if np is not None else [
        float(i) for i in range(0, int(max_hours) + 1, int(max(1, round(interval))))
    ]

    token_map: dict[str, dict[str, Any]] = {}
    for _, row in df.iterrows():
        admission_time = row.get("admission_time")
        if not isinstance(admission_time, datetime):
            admission_time = _parse_dt(admission_time)
        if not admission_time:
            continue
        payload = {"group": str(row.get(group_col) or "未分组"), "admission_time": _normalize_dt(admission_time)}
        for key in ("patient_id", "hisPid", "hisPID", "hisPatientId", "patientId", "patientID", "pid", "mrn", "hisMrn"):
            token = str(row.get(key) or "").strip()
            if token:
                token_map[token] = payload

    if not token_map:
        return {"indicators": {}}

    all_tokens = list(token_map.keys())
    admission_times = [item["admission_time"] for item in token_map.values()]
    start_min = min(admission_times)
    end_max = max(admission_times) + timedelta(hours=max_hours + interval)

    results: dict[str, Any] = {}
    for indicator in indicators or []:
        docs: list[dict[str, Any]] = []
        indicator_key = str(indicator or "").strip()
        lab_key = _trend_indicator_lab_key(indicator_key)
        source = "lab" if lab_key else "bedside"
        if source == "lab":
            aliases = LAB_NAME_MAP.get(lab_key or "", [indicator_key])
            patterns = [re.compile(f"^{re.escape(alias)}$", re.IGNORECASE) for alias in aliases]
            async for doc in db.dc_col("VI_ICU_EXAM_ITEM").find(
                {
                    "hisPid": {"$in": all_tokens},
                    "itemName": {"$in": patterns},
                },
                {"hisPid": 1, "itemName": 1, "result": 1, "itemValue": 1, "value": 1, "authTime": 1, "collectTime": 1, "testTime": 1, "reportTime": 1},
            ):
                docs.append(doc)
        else:
            codes = _indicator_codes(indicator_key)
            query = {
                "pid": {"$in": all_tokens},
                "code": {"$in": codes},
                "$or": [
                    {"time": {"$gte": start_min, "$lte": end_max}},
                    {"recordTime": {"$gte": start_min, "$lte": end_max}},
                ],
            }
            async for doc in db.col("bedside").find(query, {"pid": 1, "code": 1, "time": 1, "recordTime": 1, "fVal": 1, "intVal": 1, "strVal": 1, "value": 1}):
                docs.append(doc)

        grouped_bins: dict[str, dict[float, list[float]]] = {
            group: {hour: [] for hour in timeline} for group in groups
        }
        stat_rows: list[dict[str, Any]] = []
        for doc in docs:
            token = str(doc.get("hisPid") if source == "lab" else doc.get("pid") or "").strip()
            mapped = token_map.get(token)
            if not mapped:
                continue
            group = str(mapped["group"])
            obs_time = _extract_lab_time(doc) if source == "lab" else _parse_dt(doc.get("time") or doc.get("recordTime"))
            if not obs_time:
                continue
            if source == "lab" and (obs_time < start_min or obs_time > end_max):
                continue
            value = _extract_lab_value(doc) if source == "lab" else _extract_bedside_value(doc)
            if value is None:
                continue
            hour = (obs_time - mapped["admission_time"]).total_seconds() / 3600.0
            if hour < 0 or hour > max_hours:
                continue
            bin_hour = math.floor(hour / interval) * interval
            bin_hour = round(min(max_hours, max(0.0, bin_hour)), 3)
            if group in grouped_bins and bin_hour in grouped_bins[group]:
                grouped_bins[group][bin_hour].append(float(value))
                stat_rows.append({"group": group, "hour": bin_hour, "value": float(value)})

        group_payload: dict[str, Any] = {}
        for group in groups:
            means: list[float | None] = []
            stds: list[float | None] = []
            lowers: list[float | None] = []
            uppers: list[float | None] = []
            counts: list[int] = []
            for hour in timeline:
                arr = grouped_bins[group][hour]
                counts.append(int(len(arr)))
                if not arr:
                    means.append(None)
                    stds.append(None)
                    lowers.append(None)
                    uppers.append(None)
                    continue
                values = np.asarray(arr, dtype=float) if np is not None else arr
                if np is not None:
                    mean = float(np.mean(values))
                    std = float(np.std(values, ddof=1)) if len(values) > 1 else 0.0
                else:
                    mean = float(sum(values) / len(values))
                    std = 0.0
                se = std / math.sqrt(max(1, len(arr)))
                means.append(mean)
                stds.append(std)
                lowers.append(mean - 1.96 * se)
                uppers.append(mean + 1.96 * se)
            group_payload[group] = {
                "mean": means,
                "std": stds,
                "ci_lower": lowers,
                "ci_upper": uppers,
                "n": counts,
            }

        repeated_measures_p = None
        if smf is not None and anova_lm is not None and len(stat_rows) >= 20 and len(groups) >= 2:
            try:
                stat_df = pd.DataFrame(stat_rows)
                model = smf.ols("value ~ C(group) * C(hour)", data=stat_df).fit()
                table = anova_lm(model, typ=2)
                if "C(group):C(hour)" in table.index:
                    repeated_measures_p = float(table.loc["C(group):C(hour)", "PR(>F)"])
            except Exception:
                repeated_measures_p = None

        timepoint_comparisons: list[dict[str, Any]] = []
        if len(groups) >= 2:
            g1, g2 = groups[0], groups[1]
            for hour in timeline:
                arr1 = grouped_bins[g1][hour]
                arr2 = grouped_bins[g2][hour]
                if len(arr1) < 2 or len(arr2) < 2:
                    continue
                try:
                    _, p = stats.mannwhitneyu(arr1, arr2, alternative="two-sided")
                    timepoint_comparisons.append({"hour": hour, "p": float(p), "p_display": _p_display(float(p))})
                except Exception:
                    continue

        results[indicator] = {
            "timeline_hours": timeline,
            "source": source,
            "groups": group_payload,
            "repeated_measures_p": repeated_measures_p,
            "repeated_measures_p_display": _p_display(repeated_measures_p),
            "timepoint_comparisons": timepoint_comparisons,
        }

    return {
        "indicators": results,
        "group_by": group_by,
        "time_reference": time_reference,
        "time_range_hours": time_range_hours,
        "interval_hours": interval,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }


def _cm_to_inch(cm: float) -> float:
    return cm / 2.54


def _configure_matplotlib(config=None) -> tuple[int, str, str]:
    cfg = _get_research_cfg(config)
    dpi = int(cfg.get("figure_dpi", 300) or 300)
    font_family = str(cfg.get("figure_font_family", "Times New Roman") or "Times New Roman")
    font_cjk = str(cfg.get("figure_font_family_cjk", "SimHei") or "SimHei")
    if matplotlib is not None:
        matplotlib.rcParams["font.family"] = [font_family, font_cjk, "Arial Unicode MS", "sans-serif"]
        matplotlib.rcParams["axes.unicode_minus"] = False
    return dpi, font_family, font_cjk


def _new_figure(width_mode: str = "single", *, config=None) -> tuple[Any, Any, int]:
    if plt is None:
        raise RuntimeError("缺少 matplotlib 依赖，无法导出图表")
    dpi, _, _ = _configure_matplotlib(config)
    width_cm = FIGURE_WIDTH_CM_DOUBLE if str(width_mode).lower() == "double" else FIGURE_WIDTH_CM_SINGLE
    fig, ax = plt.subplots(figsize=(_cm_to_inch(width_cm), _cm_to_inch(width_cm * 0.72)), dpi=dpi)
    fig.patch.set_facecolor("white")
    ax.set_facecolor("white")
    return fig, ax, dpi


def _render_survival_figure(ax, data: dict[str, Any]) -> None:
    km = (data.get("kaplan_meier") or {}) if isinstance(data, dict) else {}
    curves = km.get("curves") or {}
    for name, curve in curves.items():
        timeline = curve.get("timeline") or []
        survival = curve.get("survival") or []
        lower = curve.get("ci_lower") or []
        upper = curve.get("ci_upper") or []
        ax.step(timeline, survival, where="post", linewidth=2, label=str(name))
        if lower and upper and len(lower) == len(timeline):
            ax.fill_between(timeline, lower, upper, step="post", alpha=0.18)
    ax.set_xlabel("Time (days)")
    ax.set_ylabel("Survival probability")
    ax.set_ylim(0, 1.02)
    ax.grid(alpha=0.2, linestyle="--")
    p_text = _p_display(_safe_float(km.get("log_rank_p")))
    ax.set_title(f"Kaplan-Meier Curve (Log-rank P={p_text})")
    ax.legend(frameon=False, fontsize=8)


def _render_roc_figure(ax, data: dict[str, Any]) -> None:
    curves = data.get("curves") or {}
    for name, curve in curves.items():
        fpr = curve.get("fpr") or []
        tpr = curve.get("tpr") or []
        auc = _safe_float(curve.get("auc"))
        ax.plot(fpr, tpr, linewidth=2, label=f"{name} (AUC={auc:.3f})" if auc is not None else str(name))
    ax.plot([0, 1], [0, 1], linestyle="--", color="#666666", linewidth=1)
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.set_xlabel("1 - Specificity")
    ax.set_ylabel("Sensitivity")
    ax.set_title("ROC Curves")
    ax.grid(alpha=0.2, linestyle="--")
    ax.legend(frameon=False, fontsize=8, loc="lower right")


def _render_forest_figure(ax, data: dict[str, Any], *, preferred: str = "multivariate") -> None:
    rows = data.get(preferred) or data.get("results") or data.get("univariate") or []
    if not rows:
        ax.text(0.5, 0.5, "No data", ha="center", va="center", transform=ax.transAxes)
        return
    labels = [str(row.get("variable") or row.get("subgroup") or "") for row in rows]
    est = [_safe_float(row.get("estimate")) for row in rows]
    ci_l = [_safe_float(row.get("ci_lower")) for row in rows]
    ci_u = [_safe_float(row.get("ci_upper")) for row in rows]
    y_pos = list(range(len(labels)))[::-1]
    for idx, y in enumerate(y_pos):
        if est[idx] is None or ci_l[idx] is None or ci_u[idx] is None:
            continue
        ax.plot([ci_l[idx], ci_u[idx]], [y, y], color="#2b6cb0", linewidth=1.5)
        ax.scatter([est[idx]], [y], color="#1a365d", s=20, zorder=3)
    ax.axvline(1.0, color="#999999", linestyle="--", linewidth=1)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("Effect size (OR/HR)")
    ax.set_title("Forest Plot")
    ax.grid(axis="x", alpha=0.2, linestyle="--")


def _render_trend_figure(ax, data: dict[str, Any]) -> None:
    indicators = data.get("indicators") or {}
    if not indicators:
        ax.text(0.5, 0.5, "No data", ha="center", va="center", transform=ax.transAxes)
        return
    first_key = list(indicators.keys())[0]
    item = indicators[first_key]
    timeline = item.get("timeline_hours") or []
    for group_name, values in (item.get("groups") or {}).items():
        mean = values.get("mean") or []
        lower = values.get("ci_lower") or []
        upper = values.get("ci_upper") or []
        ax.plot(timeline, mean, linewidth=2, label=group_name)
        if lower and upper and len(lower) == len(timeline):
            ax.fill_between(timeline, lower, upper, alpha=0.16)
    ax.set_xlabel("Time since ICU admission (hours)")
    ax.set_ylabel(first_key)
    ax.set_title(f"Trend Analysis - {first_key}")
    ax.grid(alpha=0.2, linestyle="--")
    ax.legend(frameon=False, fontsize=8)


def _render_correlation_heatmap(ax, data: dict[str, Any]) -> None:
    matrix = (data.get("matrix") or {}) if isinstance(data, dict) else {}
    labels = matrix.get("labels") or []
    corr = matrix.get("correlations") or []
    if np is None or not labels or not corr:
        ax.text(0.5, 0.5, "No data", ha="center", va="center", transform=ax.transAxes)
        return
    arr = np.asarray(corr, dtype=float)
    if sns is not None:
        sns.heatmap(arr, ax=ax, cmap="RdBu_r", vmin=-1, vmax=1, center=0, xticklabels=labels, yticklabels=labels, annot=True, fmt=".2f")
    else:
        im = ax.imshow(arr, cmap="RdBu_r", vmin=-1, vmax=1)
        ax.figure.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
        ax.set_xticks(range(len(labels)))
        ax.set_yticks(range(len(labels)))
        ax.set_xticklabels(labels, rotation=45, ha="right")
        ax.set_yticklabels(labels)
    ax.set_title("Correlation Heatmap")


async def export_figure(
    chart_type: str,
    result: dict[str, Any],
    fmt: str = "png",
    width_mode: str = "single",
    db=None,
    *,
    config=None,
    filename: str | None = None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    fmt_value = str(fmt or "png").lower()
    if fmt_value not in {"png", "svg", "pdf"}:
        raise ValueError("仅支持 png/svg/pdf")
    fig, ax, dpi = _new_figure(width_mode, config=config)
    chart = str(chart_type or "").lower()
    if chart in {"survival", "kaplan_meier", "km"}:
        _render_survival_figure(ax, result)
    elif chart in {"roc"}:
        _render_roc_figure(ax, result)
    elif chart in {"regression", "forest", "subgroup"}:
        _render_forest_figure(ax, result)
    elif chart in {"trend"}:
        _render_trend_figure(ax, result)
    elif chart in {"correlation", "heatmap"}:
        _render_correlation_heatmap(ax, result)
    else:
        ax.text(0.5, 0.5, "Unsupported chart type", ha="center", va="center", transform=ax.transAxes)
        ax.axis("off")

    fig.tight_layout()
    out_name = str(filename or f"{chart}_{uuid.uuid4().hex[:8]}.{fmt_value}")
    if not out_name.lower().endswith(f".{fmt_value}"):
        out_name = f"{out_name}.{fmt_value}"
    file_path = RESEARCH_EXPORT_DIR / out_name
    fig.savefig(file_path, format=fmt_value, dpi=dpi if fmt_value == "png" else None, bbox_inches="tight")
    plt.close(fig)
    return {
        "file_path": str(file_path),
        "file_name": out_name,
        "format": fmt_value,
        "dpi": dpi,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }


def _table_rows_from_payload(payload: dict[str, Any]) -> tuple[list[str], list[list[str]]]:
    groups = [str(x) for x in (payload.get("groups") or [])]
    rows = payload.get("rows") or []
    if not isinstance(rows, list):
        rows = []
    headers = ["Variable", *groups, "Statistic", "P value"]
    body: list[list[str]] = []
    for row in rows:
        if not isinstance(row, dict):
            continue
        values = row.get("values") if isinstance(row.get("values"), list) else []
        value_cells = [str(v) for v in values]
        while len(value_cells) < len(groups):
            value_cells.append("—")
        body.append([
            str(row.get("variable") or row.get("field") or ""),
            *value_cells[: len(groups)],
            str(row.get("statistic") or "—"),
            str(row.get("p_display") or _p_display(_safe_float(row.get("p_value")))),
        ])
    return headers, body


async def export_table(
    table_data: dict[str, Any],
    title: str = "Table",
    fmt: str = "docx",
    *,
    config=None,
    filename: str | None = None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    fmt_value = str(fmt or "docx").lower()
    if fmt_value not in {"docx", "csv"}:
        raise ValueError("仅支持 docx/csv")
    out_base = str(filename or f"table_{uuid.uuid4().hex[:8]}")
    headers, body = _table_rows_from_payload(table_data)

    if fmt_value == "csv":
        path = RESEARCH_EXPORT_DIR / f"{out_base}.csv"
        df = pd.DataFrame(body, columns=headers)
        df.to_csv(path, index=False, encoding="utf-8-sig")
        return {"file_path": str(path), "file_name": path.name, "format": "csv"}

    if Document is None:
        raise RuntimeError("缺少 python-docx 依赖，无法导出 Word 文档")
    path = RESEARCH_EXPORT_DIR / f"{out_base}.docx"
    doc = Document()
    doc.add_heading(str(title or table_data.get("title") or "Table"), level=1)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for row in body:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
    footnote = str(table_data.get("footnote") or "").strip()
    if footnote:
        doc.add_paragraph(footnote)
    doc.save(path)
    return {"file_path": str(path), "file_name": path.name, "format": "docx"}


async def create_materials_bundle(
    *,
    bundle_name: str,
    files: list[dict[str, Any]],
) -> dict[str, Any]:
    _require_research_analytics_deps()
    timestamp = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    safe_name = re.sub(r"[^A-Za-z0-9_\-]+", "_", str(bundle_name or "research_materials")).strip("_")
    zip_path = RESEARCH_EXPORT_DIR / f"{safe_name}_{timestamp}.zip"
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
        for item in files or []:
            file_path = str(item.get("file_path") or "").strip()
            arcname = str(item.get("arcname") or Path(file_path).name or "").strip()
            text_content = item.get("content")
            if not arcname:
                continue
            if isinstance(text_content, str):
                zf.writestr(arcname, text_content.encode("utf-8"))
                continue
            if not file_path:
                continue
            path = Path(file_path)
            if not path.exists() or not path.is_file():
                continue
            zf.write(path, arcname)
    return {"file_path": str(zip_path), "file_name": zip_path.name, "format": "zip"}


async def save_analysis_session(
    *,
    db,
    user_id: str,
    name: str,
    payload: dict[str, Any],
) -> dict[str, Any]:
    session_id = str(uuid.uuid4())
    doc = {
        "session_id": session_id,
        "name": str(name or "未命名分析会话"),
        "created_by": str(user_id or "anonymous"),
        "payload": payload,
        "created_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc),
    }
    await db.col("research_analysis_sessions").insert_one(doc)
    return {"session_id": session_id, "name": doc["name"], "created_at": serialize_doc(doc["created_at"])}


async def list_analysis_sessions(*, db, user_id: str, limit: int = 50) -> list[dict[str, Any]]:
    cursor = db.col("research_analysis_sessions").find(
        {"created_by": str(user_id or "anonymous")},
        {"_id": 0},
    ).sort("updated_at", -1).limit(max(1, min(200, int(limit or 50))))
    return [serialize_doc(doc) async for doc in cursor]


async def get_analysis_session(*, db, session_id: str, user_id: str) -> dict[str, Any] | None:
    doc = await db.col("research_analysis_sessions").find_one(
        {"session_id": str(session_id), "created_by": str(user_id or "anonymous")},
        {"_id": 0},
    )
    return serialize_doc(doc) if doc else None


def _score_value_from_doc(doc: dict[str, Any]) -> float | None:
    for key in ("score", "value", "score_value", "total", "result"):
        val = _as_number(doc.get(key))
        if val is not None:
            return val
    summary = doc.get("summary")
    if isinstance(summary, dict):
        for key in ("score", "total", "value"):
            val = _as_number(summary.get(key))
            if val is not None:
                return val
    if isinstance(summary, (int, float, str)):
        return _as_number(summary)
    return None


async def _attach_score_extrema(df: pd.DataFrame, db, score_type: str | list[str], column_name: str) -> None:
    if df.empty:
        df[column_name] = None
        return
    ids = [str(pid).strip() for pid in df.get("patient_id", pd.Series([], dtype=str)).dropna().tolist() if str(pid).strip()]
    if not ids:
        df[column_name] = None
        return
    score_types = [str(item).strip() for item in (score_type if isinstance(score_type, list) else [score_type]) if str(item or '').strip()]
    query = {"pid": {"$in": ids}, "scoreType": {"$in": score_types}}
    cursor = db.col("score").find(
        query,
        {"pid": 1, "patient_id": 1, "score": 1, "scoreType": 1, "total": 1, "value": 1, "score_value": 1, "result": 1, "summary": 1},
    )
    values: dict[str, float] = {}
    async for doc in cursor:
        pid = str(doc.get("pid") or doc.get("patient_id") or "").strip()
        if not pid:
            continue
        score = _score_value_from_doc(doc)
        if score is None:
            continue
        prev = values.get(pid)
        if prev is None or score > prev:
            values[pid] = score
    df[column_name] = df.get("patient_id", pd.Series([], dtype=str)).map(lambda pid: values.get(str(pid).strip()))


async def _attach_lab_admission_values(
    df: pd.DataFrame,
    db,
    lab_map: dict[str, list[str]],
) -> None:
    """Attach first-within-24h lab values after ICU admission for each patient."""
    if df.empty:
        for col in [f"{k}_admission" for k in lab_map]:
            df[col] = None
        return
    ids = [str(pid).strip() for pid in df.get("hisPid", pd.Series([], dtype=str)).dropna().tolist() if str(pid).strip()]
    if not ids:
        for col in [f"{k}_admission" for k in lab_map]:
            df[col] = None
        return

    # Build flat alias -> canonical key lookup
    import re
    alias_to_key: dict[str, str] = {}
    valid_name_regexes = []
    
    for key, aliases in lab_map.items():
        for alias in aliases:
            a = alias.lower()
            if a not in alias_to_key:
                alias_to_key[a] = key
                valid_name_regexes.append(re.compile(f"^{re.escape(alias)}$", re.IGNORECASE))

    # Build admission_time lookup per patient (keyed by hisPid)
    admission_map: dict[str, datetime | None] = {}
    for _, row in df.iterrows():
        pid = str(row.get("hisPid") or "").strip()
        if not pid: continue
        admission_map[pid] = row.get("admission_time")

    results: dict[str, dict[str, float]] = {pid: {} for pid in ids}
    cursor = db.dc_col("VI_ICU_EXAM_ITEM").find(
        {
            "hisPid": {"$in": ids},
            "itemName": {"$in": valid_name_regexes}
        },
        {"hisPid": 1, "itemName": 1, "result": 1, "itemValue": 1, "authTime": 1, "testTime": 1, "reportTime": 1},
    )
    async for doc in cursor:
        pid = str(doc.get("hisPid") or "").strip()
        if pid not in results:
            continue
        item_name = str(doc.get("itemName") or "").strip().lower()
        key = alias_to_key.get(item_name)
        if key is None:
            continue
        if key in results[pid]:  # already got first value
            continue
        admission_time = admission_map.get(pid)
        if admission_time is not None:
            record_time = _parse_dt(doc.get("authTime") or doc.get("testTime") or doc.get("reportTime"))
            if record_time is None:
                continue
            delta_hours = (record_time - admission_time).total_seconds() / 3600
            if not (0 <= delta_hours <= 24):
                continue
        value = _as_number(doc.get("result") or doc.get("itemValue"))
        if value is not None:
            results[pid][key] = value

    pid_series = df.get("patient_id", pd.Series([], dtype=str))
    for key in lab_map:
        col = f"{key}_admission"
        df[col] = df.get("hisPid", pd.Series([], dtype=str)).astype(str).map(lambda hpid, k=key: results.get(str(hpid).strip(), {}).get(k))


async def _attach_treatment_days(df: pd.DataFrame, db) -> None:
    """Attach vasopressor_days, mv_days, and icu_readmission per patient."""
    if df.empty:
        df["vasopressor_days"] = None
        df["mv_days"] = None
        df["icu_readmission"] = None
        return
    ids = [str(pid).strip() for pid in df.get("patient_id", pd.Series([], dtype=str)).dropna().tolist() if str(pid).strip()]
    if not ids:
        df["vasopressor_days"] = None
        df["mv_days"] = None
        df["icu_readmission"] = None
        return

    vaso_days: dict[str, float] = {}
    mv_days: dict[str, float] = {}
    
    # 1. Primary source: score collection (already calculated summaries)
    cursor = db.col("score").find(
        {"pid": {"$in": ids}, "scoreType": {"$in": ["vasopressor_days", "mv_days", "vasopressor_day", "mv_day"]}},
        {"pid": 1, "scoreType": 1, "score": 1, "value": 1, "score_value": 1, "total": 1},
    )
    async for doc in cursor:
        pid = str(doc.get("pid") or doc.get("patient_id") or "").strip()
        stype = str(doc.get("scoreType") or doc.get("score_type") or "")
        val = _score_value_from_doc(doc)
        if val is None:
            continue
        if "vasopressor" in stype:
            vaso_days[pid] = max(vaso_days.get(pid, 0), val)
        elif "mv" in stype:
            mv_days[pid] = max(mv_days.get(pid, 0), val)

    # 2. Fallback: calculate from deviceBind and drugExe for missing values
    missing_mv_ids = [pid for pid in ids if pid not in mv_days]
    if missing_mv_ids:
        # deviceBind for ventilators
        device_cursor = db.col("deviceBind").find(
            {"pid": {"$in": missing_mv_ids}},
            {"pid": 1, "bindTime": 1, "unBindTime": 1, "type": 1, "deviceName": 1}
        )
        now = datetime.now(timezone.utc)
        async for doc in device_cursor:
            pid = str(doc.get("pid") or "").strip()
            d_type = str(doc.get("type") or "").lower()
            d_name = str(doc.get("deviceName") or "").lower()
            if not any(k in d_type or k in d_name for k in ["vent", "呼吸机", "创呼吸", "无创"]):
                continue
            start = _parse_dt(doc.get("bindTime"))
            if not start: continue
            end = _parse_dt(doc.get("unBindTime")) or now
            duration_days = max(0, (end - start).total_seconds() / 86400.0)
            mv_days[pid] = mv_days.get(pid, 0) + duration_days

    missing_vaso_ids = [pid for pid in ids if pid not in vaso_days]
    if missing_vaso_ids:
        # drug records from DataCenter orders
        hpid_to_pid = {str(row.get("hisPid")).strip(): str(row.get("patient_id")).strip() for _, row in df.iterrows() if row.get("hisPid")}
        hpids = list(hpid_to_pid.keys())
        if hpids:
            vaso_keywords = ["去甲", "多巴", "肾上腺", "加压素", "norepinephrine", "epinephrine", "vasopressin", "dopamine"]
            pattern = "|".join(vaso_keywords)
            order_cursor = db.dc_col("VI_ICU_ZYYZ").find(
                {"pid": {"$in": hpids}, "orderName": {"$regex": pattern, "$options": "i"}},
                {"pid": 1, "startTime": 1, "endTime": 1, "stopTime": 1, "orderTime": 1}
            )
            async for doc in order_cursor:
                hpid = str(doc.get("pid") or "").strip()
                pid = hpid_to_pid.get(hpid)
                if not pid: continue
                # Calculate duration from order span
                start = _parse_dt(doc.get("startTime") or doc.get("orderTime"))
                if not start: continue
                end = _parse_dt(doc.get("stopTime") or doc.get("endTime")) or datetime.now(timezone.utc)
                duration_days = max(0.1, (end - start).total_seconds() / 86400.0)
                vaso_days[pid] = vaso_days.get(pid, 0) + duration_days

    pid_series = df.get("patient_id", pd.Series([], dtype=str))
    df["vasopressor_days"] = pid_series.map(lambda pid: vaso_days.get(str(pid).strip()))
    df["mv_days"] = pid_series.map(lambda pid: mv_days.get(str(pid).strip()))

    # icu_readmission: infer from multiple patient records with same hisPid
    his_pid_series = df.get("hisPid", pd.Series([], dtype=str)).astype(str)
    his_pid_counts = his_pid_series.value_counts()
    df["icu_readmission"] = his_pid_series.map(lambda hpid: 1 if his_pid_counts.get(str(hpid).strip(), 0) > 1 else 0)


async def _attach_alert_tags(df: pd.DataFrame, db) -> None:
    if df.empty:
        df["alert_tags"] = []
        return
    ids = [str(pid).strip() for pid in df.get("patient_id", pd.Series([], dtype=str)).dropna().tolist() if str(pid).strip()]
    if not ids:
        df["alert_tags"] = [[] for _ in range(len(df))]
        return
    tags_map: dict[str, set[str]] = {}
    cursor = db.col("alert_records").find(
        {"patient_id": {"$in": ids}},
        {"patient_id": 1, "alert_type": 1, "scanner": 1, "rule_id": 1, "rule_name": 1, "extra": 1},
    )
    async for doc in cursor:
        pid = str(doc.get("patient_id") or "").strip()
        if not pid:
            continue
        tags = tags_map.setdefault(pid, set())
        for field in ("alert_type", "scanner", "rule_id", "rule_name"):
            text = str(doc.get(field) or "").strip()
            if text:
                tags.add(text.lower())
                tags.add(text)
        extra = doc.get("extra") if isinstance(doc.get("extra"), dict) else {}
        if isinstance(extra, dict):
            for field in ("scanner", "scanner_code", "category", "alert_code"):
                text = str(extra.get(field) or "").strip()
                if text:
                    tags.add(text.lower())
                    tags.add(text)
    df["alert_tags"] = df.get("patient_id", pd.Series([], dtype=str)).map(lambda pid: sorted(tags_map.get(str(pid).strip(), set())))


def _normalize_bool_flag(value: Any) -> bool | None:
    if isinstance(value, bool):
        return value
    text = str(value or "").strip().lower()
    if text in {"1", "true", "yes", "y", "是", "有"}:
        return True
    if text in {"0", "false", "no", "n", "否", "无"}:
        return False
    return None


def _normalize_sex_token(value: Any) -> str | None:
    text = str(value or "").strip().upper()
    if text in {"M", "MALE", "男"}:
        return "M"
    if text in {"F", "FEMALE", "女"}:
        return "F"
    return None


def _normalize_outcome_filter(value: Any) -> str | None:
    text = str(value or "").strip().lower()
    if not text:
        return None
    if any(word in text for word in ["dead", "death", "亡", "死亡", "icu内死亡", "icu死亡"]):
        return "dead"
    if any(word in text for word in ["alive", "存活", "出科", "转科", "出院"]):
        return "alive"
    return None


def _builder_field_alias(field: str) -> str:
    mapping = {
        "diagnosis": "primary_diagnosis",
        "diagnosis_text": "primary_diagnosis",
        "age": "age",
        "sex": "sex",
        "icu_los": "los_icu_days",
        "icu_days": "los_icu_days",
        "los_icu_days": "los_icu_days",
        "sofa": "sofa_max",
        "sofa_max": "sofa_max",
        "apache2": "apache2_max",
        "apache_ii": "apache2_max",
        "apache2_max": "apache2_max",
        "mechanical_ventilation": "mechanical_ventilation",
        "crrt": "crrt",
        "vasopressor": "vasopressor",
        "primary_diagnosis_category": "primary_category",
        "diagnosis_category": "primary_category",
        "outcome": "outcome",
        "admission_time": "admission_time",
        "alert_type": "alert_tags",
    }
    key = str(field or "").strip()
    return mapping.get(key, key)


def _empty_mask(df: pd.DataFrame) -> pd.Series:
    return pd.Series([True] * len(df), index=df.index)


def _value_range(value: Any) -> tuple[float | None, float | None]:
    if isinstance(value, (list, tuple)) and value:
        start = _as_number(value[0])
        end = _as_number(value[1]) if len(value) > 1 else None
        return start, end
    return _as_number(value), None


def _date_range(value: Any) -> tuple[datetime | None, datetime | None]:
    if isinstance(value, (list, tuple)) and value:
        start = _parse_dt(value[0])
        end = _parse_dt(value[1]) if len(value) > 1 else None
        return start, end
    if isinstance(value, dict):
        return _parse_dt(value.get("start")), _parse_dt(value.get("end"))
    return _parse_dt(value), None


def _builder_filter_mask(df: pd.DataFrame, condition: dict[str, Any]) -> pd.Series:
    if df.empty:
        return _empty_mask(df)
    try:
        field = _builder_field_alias(condition.get("field"))
        operator = str(condition.get("operator") or "eq").strip().lower()
        value = condition.get("value")
        if field == "primary_diagnosis":
            series = df.get("primary_diagnosis", pd.Series(["" for _ in range(len(df))], index=df.index)).astype(str)
            keywords: list[str] = []
            if isinstance(value, list):
                keywords = [str(item).strip() for item in value if str(item or "").strip()]
            elif isinstance(value, dict):
                keywords = [str(value.get("code") or "").strip(), str(value.get("name") or "").strip()]
                keywords = [item for item in keywords if item]
            else:
                text = str(value or "").strip()
                if text:
                    keywords = [item.strip() for item in text.split("|") if item.strip()]
            if not keywords:
                return _empty_mask(df)
            mask = pd.Series([False for _ in range(len(df))], index=df.index)
            for token in keywords:
                mask = mask | series.str.contains(token, case=False, na=False, regex=False)
            if operator == "not_contains":
                return ~mask
            return mask
        if field == "primary_category":
            series = df.get("primary_category", pd.Series(["" for _ in range(len(df))], index=df.index)).astype(str)
            text = str(value or "").strip()
            if not text:
                return _empty_mask(df)
            if operator == "contains":
                return series.str.contains(text, case=False, na=False, regex=False)
            return series.str.lower() == text.lower()
        if field == "sex":
            target = _normalize_sex_token(value)
            if not target:
                return _empty_mask(df)
            series = df.get("sex", pd.Series([None for _ in range(len(df))], index=df.index)).astype(str).str.upper()
            return series == target
        if field in {"mechanical_ventilation", "crrt", "vasopressor"}:
            target_bool = _normalize_bool_flag(value)
            if target_bool is None:
                return _empty_mask(df)
            raw = df.get(field, pd.Series([None for _ in range(len(df))], index=df.index))
            series = raw.apply(_coerce_binary).fillna(0).astype(int)
            return series == (1 if target_bool else 0)
        if field in {"age", "los_icu_days", "sofa_max", "apache2_max"}:
            series = pd.to_numeric(df.get(field, pd.Series([None for _ in range(len(df))], index=df.index)), errors="coerce")
            if operator == "range":
                low, high = _value_range(value)
                mask = _empty_mask(df)
                if low is not None:
                    mask = mask & (series >= low)
                if high is not None:
                    mask = mask & (series <= high)
                return mask
            target = _as_number(value)
            if target is None:
                return _empty_mask(df)
            if operator in {"gt", ">"}:
                return series > target
            if operator in {"gte", ">="}:
                return series >= target
            if operator in {"lt", "<"}:
                return series < target
            if operator in {"lte", "<="}:
                return series <= target
            return series == target
        if field == "outcome":
            target = _normalize_outcome_filter(value)
            if not target:
                return _empty_mask(df)
            series = df.get("outcome", pd.Series(["" for _ in range(len(df))], index=df.index)).astype(str).str.lower()
            return series == target
        if field == "admission_time":
            series = pd.to_datetime(
                df.get("admission_time", pd.Series([None for _ in range(len(df))], index=df.index)),
                errors="coerce",
                utc=True,
            )
            start, end = _date_range(value)
            mask = _empty_mask(df)
            if start:
                mask = mask & (series >= start)
            if end:
                mask = mask & (series <= end)
            return mask
        if field == "alert_tags":
            target = str(value or "").strip().lower()
            tags_series = df.get("alert_tags", pd.Series([[] for _ in range(len(df))], index=df.index))
            has_tag = tags_series.apply(lambda tags: any(str(tag).strip().lower() == target for tag in tags or []))
            if operator in {"not_exists", "not", "neq"}:
                return ~has_tag
            return has_tag
        return _empty_mask(df)
    except Exception as exc:
        logger.warning("builder filter failed field=%s operator=%s: %s", condition.get("field"), condition.get("operator"), exc)
        return _empty_mask(df)


def _apply_builder_filters(df: pd.DataFrame, filters: list[dict[str, Any]]) -> pd.DataFrame:
    if df.empty or not filters:
        return df
    mask = _empty_mask(df)
    for item in filters:
        mask = mask & _builder_filter_mask(df, item if isinstance(item, dict) else {})
    return df[mask].copy()


def _cohort_demographics_from_df(df: pd.DataFrame) -> dict[str, Any]:
    if df.empty:
        return {}
    total = len(df)
    demographics: dict[str, Any] = {"n": total}
    age_series = pd.to_numeric(df.get("age"), errors="coerce").dropna()
    if not age_series.empty:
        demographics["age_mean"] = float(age_series.mean())
        demographics["age_std"] = float(age_series.std(ddof=1)) if len(age_series) > 1 else 0.0
    los_series = pd.to_numeric(df.get("los_icu_days"), errors="coerce").dropna()
    if not los_series.empty:
        demographics["los_median"] = float(los_series.median())
        demographics["los_q1"] = float(los_series.quantile(0.25))
        demographics["los_q3"] = float(los_series.quantile(0.75))
    sofa_series = pd.to_numeric(df.get("sofa_admission"), errors="coerce").dropna()
    if not sofa_series.empty:
        demographics["sofa_mean"] = float(sofa_series.mean())
        demographics["sofa_std"] = float(sofa_series.std(ddof=1)) if len(sofa_series) > 1 else 0.0
    sex_series = df.get("sex", pd.Series([None for _ in range(len(df))], index=df.index)).astype(str).str.upper()
    demographics["male_ratio"] = float((sex_series == "M").sum()) / total if total else 0.0
    mortality_series = pd.to_numeric(df.get("icu_mortality"), errors="coerce").fillna(0)
    if not mortality_series.empty:
        demographics["mortality_rate"] = float(mortality_series.mean())
    return demographics


async def build_custom_cohort(
    *,
    db,
    filters: list[dict[str, Any]] | None = None,
    patient_ids: list[str] | None = None,
    department: str | None = None,
    dept_code: str | None = None,
    patient_scope: str = 'all',
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    scoped = _ensure_patient_id_set(patient_ids)
    if not scoped:
        dept_name = str(department or "").strip()
        dept_code_text = str(dept_code or "").strip()
        # 前端在科室名未加载时可能把科室编码误传到 department；这里自动纠偏，避免 AND 条件把队列过滤为 0。
        if dept_name and not dept_code_text and dept_name.isdigit():
            dept_code_text = dept_name
            dept_name = ""
        if dept_name and dept_code_text and (dept_name == dept_code_text or dept_name.isdigit()):
            dept_name = ""
        base_query = research_patient_scope_query(patient_scope)
        clauses = [base_query] if base_query else []
        if dept_name:
            clauses.append({"$or": [{"hisDept": dept_name}, {"dept": dept_name}]})
        if dept_code_text:
            clauses.append({"deptCode": dept_code_text})
        query = {} if not clauses else clauses[0] if len(clauses) == 1 else {"$and": clauses}
        cursor = db.col("patient").find(query, {"_id": 1}).limit(max_patients)
        scoped = []
        async for doc in cursor:
            pid = str(doc.get("_id") or "").strip()
            if pid:
                scoped.append(pid)
    scoped = _ensure_patient_id_set(scoped)
    if not scoped:
        return {"patient_ids": [], "patient_count": 0, "patient_scope": patient_scope, "demographics": {}, "preview_patients": []}
    df = await _load_patient_dataframe(scoped, db, max_patients=max_patients)
    if df.empty:
        return {"patient_ids": [], "patient_count": 0, "patient_scope": patient_scope, "demographics": {}, "preview_patients": []}
    primary_category = None
    for col in ("primaryDiagnosisCategory", "mainDiagnosisCategory", "diagnosisCategory", "diagnosis_category", "primary_category"):
        if col in df.columns:
            primary_category = df[col]
            break
    if primary_category is None:
        primary_category = pd.Series(["" for _ in range(len(df))], index=df.index)
    df["primary_category"] = primary_category
    await _attach_score_extrema(df, db, "sofa", "sofa_max")
    await _attach_score_extrema(df, db, ["apache2", "apacheII"], "apache2_max")
    await _attach_alert_tags(df, db)
    filtered = _apply_builder_filters(df, filters or [])
    patient_ids_out = filtered.get("patient_id", pd.Series([], dtype=str)).dropna().astype(str).tolist()[:max_patients]
    preview_rows: list[dict[str, Any]] = []
    for _, row in filtered.head(10).iterrows():
        preview_rows.append(
            {
                "id": str(row.get("patient_id") or row.get("hisPid") or ""),
                "age": _as_number(row.get("age")),
                "sex": "男" if str(row.get("sex") or "").upper() == "M" else "女" if str(row.get("sex") or "").upper() == "F" else "-",
                "diagnosis": row.get("primary_diagnosis") or row.get("clinicalDiagnosis") or "",
                "los_days": _as_number(row.get("los_icu_days")),
                "outcome": "死亡" if str(row.get("outcome") or "").lower() == "dead" else "存活",
            }
        )
    demographics = _cohort_demographics_from_df(filtered)
    return {
        "patient_ids": patient_ids_out,
        "patient_count": len(patient_ids_out),
        "patient_scope": patient_scope,
        "demographics": demographics,
        "preview_patients": preview_rows,
    }


def _summaries_from_series(series: pd.Series) -> dict[str, Any]:
    total = int(len(series))
    non_null = int(series.notna().sum())
    result: dict[str, Any] = {
        "total_count": total,
        "non_null_count": non_null,
        "non_null_rate": (non_null / total) if total else None,
    }
    numeric = pd.to_numeric(series, errors="coerce")
    numeric = numeric.dropna()
    if not numeric.empty:
        result["mean"] = float(numeric.mean())
        result["std"] = float(numeric.std(ddof=1)) if len(numeric) > 1 else 0.0
        result["median"] = float(numeric.median())
        result["range"] = {"min": float(numeric.min()), "max": float(numeric.max())}
    value_counts = series.dropna().astype(str).value_counts().head(8)
    if not value_counts.empty:
        distribution: dict[str, Any] = {}
        for value, count in value_counts.items():
            distribution[str(value)] = {"count": int(count), "ratio": (int(count) / total) if total else None}
        result["distribution"] = distribution
    return result


async def summarize_variables(
    *,
    db,
    patient_ids: list[str] | None = None,
    cohort_id: str | None = None,
    fields: list[str] | None = None,
    config=None,
) -> dict[str, Any]:
    _require_research_analytics_deps()
    resolved = await _resolve_patient_ids(patient_ids, cohort_id, db)
    if not resolved:
        return {"summaries": {}, "n_patients": 0}
    cfg = _get_research_cfg(config)
    max_patients = int(cfg.get("max_export_patients", 10000) or 10000)
    df = await _load_patient_dataframe(resolved, db, max_patients=max_patients)
    if df.empty:
        return {"summaries": {}, "n_patients": 0}
    summaries: dict[str, Any] = {}
    for field in fields or []:
        name = str(field or "").strip()
        if not name or name not in df.columns:
            continue
        summaries[name] = _summaries_from_series(df[name])
    return {"summaries": summaries, "n_patients": len(df)}
