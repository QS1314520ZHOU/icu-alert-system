"""
DOCX exporter for ICU clinical documents.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def set_run_font(
    run,
    font_name: str = "SimSun",
    font_size: float = 10.5,
    bold: bool = False,
    italic: bool = False,
    color_rgb: RGBColor | None = None,
):
    """Set font styling including East Asia font support in python-docx."""
    run.font.name = "Times New Roman"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color_rgb:
        run.font.color.rgb = color_rgb


def _content(draft_doc: dict) -> dict:
    return draft_doc.get("current_content") or draft_doc.get("draft") or {}


def _is_workbench(content: dict) -> bool:
    return content.get("content_type") == "rounding_workbench"


def _add_heading(doc: Document, text: str, level_size: float = 12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, font_name="Microsoft YaHei", font_size=level_size, bold=True, color_rgb=RGBColor(22, 119, 255))
    return p


def _add_text(doc: Document, text: Any, indent: float = 0.2):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(str(text or "未提供"))
    set_run_font(run, font_name="SimSun", font_size=10.5)
    return p


def _add_bullet(doc: Document, text: str, indent: float = 0.25):
    _add_text(doc, f"• {text}", indent)


def _setup_doc() -> Document:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("ICU 24小时病程记录")
    set_run_font(title_run, font_name="Microsoft YaHei", font_size=16, bold=True, color_rgb=RGBColor(22, 119, 255))
    return doc


def export_progress_note_docx(draft: dict) -> bytes:
    """Generate a DOCX file for new workbench drafts or legacy SOAP drafts."""
    content = _content(draft)
    doc = _setup_doc()
    if _is_workbench(content):
        _export_workbench(doc, draft, content)
    else:
        _export_legacy_soap(doc, draft, content)
    _add_signature(doc, draft)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _export_workbench(doc: Document, draft_doc: dict, content: dict) -> None:
    banner = content.get("patient_banner") or {}
    patient_id = draft_doc.get("patient_id") or ""

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["床号", "患者ID", "性别 / 年龄", "ICU天数"]
    values = [
        banner.get("bed_no") or "未提供",
        patient_id or "未提供",
        f"{banner.get('sex') or '未提供'} / {banner.get('age') or '未提供'}岁",
        f"第{banner.get('icu_day') or '未提供'}天",
    ]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, font_name="Microsoft YaHei", font_size=10, bold=True)
    for idx, text in enumerate(values):
        cell = table.rows[1].cells[idx]
        cell.text = str(text)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, font_name="SimSun", font_size=10)

    _add_heading(doc, "一、患者摘要")
    _add_text(doc, f"主要诊断：{banner.get('primary_diagnosis') or '未提供'}")
    _add_text(doc, f"过敏：{banner.get('allergy_status') or '未提供'}；隔离：{banner.get('isolation_status') or '未提供'}；代码状态：{banner.get('code_status') or '未提供'}")

    _add_heading(doc, "二、器官支持")
    for item in content.get("organ_support") or []:
        missing = item.get("missing_data") or []
        suffix = f"（缺失：{'、'.join(missing)}）" if missing else ""
        _add_bullet(doc, f"{item.get('label')}: {item.get('summary')}{suffix}")

    _add_heading(doc, "三、过去24小时关键事件")
    timeline = content.get("timeline") or []
    if timeline:
        for event in timeline:
            _add_bullet(doc, f"{event.get('occurred_at')} {event.get('title')}：{event.get('description')}")
    else:
        _add_text(doc, "未提供关键事件。")

    _add_heading(doc, "四、Assessment & Plan")
    for idx, card in enumerate(content.get("system_ap") or [], start=1):
        _add_text(doc, f"{idx}. {card.get('title') or card.get('system')}", indent=0.1)
        for label, key in (("状态", "status"), ("趋势", "trend"), ("评估", "assessment"), ("计划", "plan_items")):
            statements = [s.get("text") for s in card.get(key) or [] if s.get("text")]
            if statements:
                _add_text(doc, f"{label}：" + "；".join(statements), indent=0.35)

    _add_heading(doc, "五、今日目标")
    for goal in content.get("daily_goals") or []:
        mark = "☑" if goal.get("status") == "done" else "☐"
        _add_bullet(doc, f"{mark} {goal.get('label')}: {goal.get('target') or '待确认'}")

    _add_heading(doc, "六、风险任务与待复核项")
    risk_tasks = content.get("risk_tasks") or []
    if risk_tasks:
        for task in risk_tasks:
            _add_bullet(doc, f"{task.get('title')}（{task.get('priority')}）：{'；'.join(task.get('confirm_items') or [])}")
    else:
        _add_text(doc, "暂无活跃风险任务。")

    _add_heading(doc, "七、可签署病程文本")
    preview = content.get("note_preview") or {}
    final_text = preview.get("final_text_override") or preview.get("generated_text") or "未提供"
    for paragraph in str(final_text).splitlines():
        if paragraph.strip():
            _add_text(doc, paragraph, indent=0.2)


def _export_legacy_soap(doc: Document, draft_doc: dict, content: dict) -> None:
    snapshot = draft_doc.get("context_snapshot") or {}
    basics = snapshot.get("basics") or {}
    patient_id = draft_doc.get("patient_id") or ""
    patient_name = basics.get("name") or snapshot.get("patient_name") or "未提供"

    table = doc.add_table(rows=2, cols=4)
    table.style = "Table Grid"
    headers = ["床号", "患者姓名 / ID", "性别 / 年龄", "ICU天数"]
    values = [
        str(basics.get("bed") or "未提供"),
        f"{patient_name} ({patient_id})",
        f"{basics.get('sex') or '未提供'} / {basics.get('age') or '未提供'}岁",
        f"第{basics.get('day') or '未提供'}天",
    ]
    for idx, text in enumerate(headers):
        table.rows[0].cells[idx].text = text
    for idx, text in enumerate(values):
        table.rows[1].cells[idx].text = text

    _add_heading(doc, "主要诊断")
    _add_text(doc, basics.get("diagnosis") or "未提供")

    def add_soap_section(tag: str, title_text: str, data: Any):
        _add_heading(doc, f"【{tag}】{title_text}")
        if isinstance(data, str):
            _add_text(doc, data)
        elif isinstance(data, dict):
            labels = {
                "vitals": "生命体征",
                "labs": "检验",
                "drugs": "用药",
                "ventilator": "呼吸机",
                "alerts": "预警",
            }
            for sub_key, sub_val in data.items():
                if sub_val:
                    _add_text(doc, f"{labels.get(sub_key, sub_key)}：{sub_val}")
        elif isinstance(data, list):
            for idx, item in enumerate(data, start=1):
                if item:
                    _add_text(doc, f"{idx}. {item}")

    add_soap_section("S", "主观症状", content.get("subjective") or "未提供")
    add_soap_section("O", "客观资料", content.get("objective") or {})
    add_soap_section("A", "病情评估", content.get("assessment") or {})
    add_soap_section("P", "诊疗计划", content.get("plan") or [])


def _add_signature(doc: Document, draft: dict) -> None:
    if draft.get("status") != "finalized":
        return
    sig_p = doc.add_paragraph()
    sig_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    signer = draft.get("finalized_by") or "系统"
    sig_date = draft.get("finalized_at")
    if isinstance(sig_date, str):
        try:
            sig_date = datetime.fromisoformat(sig_date)
        except ValueError:
            pass
    sig_str = sig_date.strftime("%Y-%m-%d %H:%M") if isinstance(sig_date, datetime) else str(sig_date or "")
    sig_run = sig_p.add_run(f"医师签名：{signer}      签署时间：{sig_str}")
    set_run_font(sig_run, font_name="SimSun", font_size=11, bold=True, italic=True)
